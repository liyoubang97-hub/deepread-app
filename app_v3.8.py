"""
DeepRead V3.8 - 导出功能版
改进：
1. 添加Markdown导出功能
2. 支持导出个人笔记
3. 支持导出完整学习笔记
4. 可导入飞书文档
"""

import streamlit as st
from pathlib import Path
import sys
import time
from datetime import datetime, timedelta
from io import BytesIO
from PIL import Image, ImageDraw, ImageFont
import base64
import matplotlib
matplotlib.use('Agg')  # 非交互式后端
import matplotlib.pyplot as plt
import matplotlib.patches as patches
from matplotlib import font_manager

# PDF和Word导出
try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib import colors
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_SUPPORT = True
except ImportError:
    WORD_SUPPORT = False

sys.path.insert(0, str(Path(__file__).parent))

from lazy_loader import get_book_content, get_cache_info, clear_cache
from practice_tasks_enhanced import PRACTICE_TASKS

st.set_page_config(
    page_title="DeepRead 深读",
    page_icon="🧠",
    layout="centered",
    initial_sidebar_state="expanded"
)

# 优化的样式
st.markdown("""
<style>
    /* 字体引入 */
    @import url('https://fonts.googleapis.com/css2?family=Noto+Serif+SC:wght@400;500;600;700&family=Inter:wght@400;500;600&display=swap');

    /* 主容器 */
    .main {
        padding: 0 !important;
        max-width: 900px !important;
        margin: 0 auto;
        background: #FFFFFF;
    }

    /* 隐藏默认元素 */
    #MainMenu, footer, .stDeployButton {
        visibility: hidden;
        display: none !important;
    }

    /* ===== 页面头部 ===== */
    .page-header {
        text-align: center;
        padding: 2.5rem 0 2rem 0;
        border-bottom: 2px solid #E8EEF2;
        margin-bottom: 2rem;
    }

    .page-title {
        font-family: 'Noto Serif SC', serif;
        font-size: 2rem;
        font-weight: 700;
        color: #2D3436;
        letter-spacing: 0.02em;
        margin-bottom: 0.5rem;
    }

    .page-subtitle {
        font-family: 'Inter', sans-serif;
        font-size: 0.9rem;
        font-weight: 400;
        color: #636E72;
        letter-spacing: 0.05em;
        text-transform: uppercase;
    }

    /* ===== 分区标题 ===== */
    .section-block {
        margin-bottom: 2.5rem;
    }

    .section-title {
        font-family: 'Noto Serif SC', serif;
        font-size: 1.5rem;
        font-weight: 600;
        color: #2D3436;
        margin-bottom: 1.25rem;
        padding-bottom: 0.5rem;
        border-bottom: 2px solid #2D3436;
        display: inline-block;
    }

    .section-subtitle {
        font-family: 'Inter', sans-serif;
        font-size: 0.95rem;
        font-weight: 500;
        color: #636E72;
        margin-bottom: 1.25rem;
        font-style: italic;
    }

    /* ===== 内容卡片 - 极简设计 ===== */
    .content-block {
        background: #FFFFFF;
        border-radius: 0;
        padding: 1.5rem 0;
        margin: 1rem 0;
        border: none;
        border-left: 3px solid #E8EEF2;
    }

    .content-block.highlight {
        border-left: 3px solid #2D3436;
    }

    /* ===== 核心观点 - 移除米色背景，改为极简线条 ===== */
    .core-idea-box {
        background: #FFFFFF;
        color: #2D3436;
        padding: 1.5rem 0;
        margin: 1.5rem 0;
        border-left: 4px solid #2D3436;
        border-radius: 0;
    }

    .core-idea-text {
        font-family: 'Noto Serif SC', serif;
        font-size: 1.05rem;
        font-weight: 500;
        line-height: 1.8;
        white-space: pre-wrap;
    }

    /* ===== 小标题 ===== */
    .subsection-header {
        font-family: 'Noto Serif SC', serif;
        font-size: 1.1rem;
        font-weight: 600;
        color: #2D3436;
        margin-top: 1.5rem;
        margin-bottom: 0.75rem;
    }

    /* ===== 正文 ===== */
    .body-text {
        font-family: 'Noto Serif SC', serif;
        font-size: 1rem;
        line-height: 1.85;
        color: #2D3436;
        margin-bottom: 0.75rem;
    }

    /* ===== 提问框 - 移除背景色 ===== */
    .question-block {
        background: #FFFFFF;
        border-left: 3px solid #636E72;
        border-radius: 0;
        padding: 1.5rem 0;
        margin: 1.5rem 0;
    }

    .question-label {
        font-family: 'Inter', sans-serif;
        font-size: 0.85rem;
        font-weight: 600;
        color: #636E72;
        font-style: italic;
        letter-spacing: 0.02em;
        margin-bottom: 0.5rem;
    }

    .question-text {
        font-family: 'Noto Serif SC', serif;
        font-size: 1.05rem;
        font-weight: 500;
        color: #2D3436;
        line-height: 1.7;
    }

    .hint-text {
        font-family: 'Inter', sans-serif;
        font-size: 0.875rem;
        color: #636E72;
        margin-top: 0.75rem;
        font-style: italic;
    }

    /* ===== 金句卡片 - 移除背景色 ===== */
    .quote-block {
        font-family: 'Noto Serif SC', serif;
        font-size: 1.05rem;
        font-style: italic;
        color: #2D3436;
        line-height: 1.75;
        padding: 1.5rem 0;
        background: #FFFFFF;
        border-radius: 0;
        margin: 1.25rem 0;
        position: relative;
        border-left: 3px solid #636E72;
    }

    .quote-block::before {
        content: '"';
        font-size: 3rem;
        color: #636E72;
        opacity: 0.15;
        position: absolute;
        top: -0.25rem;
        left: 1.25rem;
        font-family: Georgia, serif;
    }

    /* ===== 分隔线 ===== */
    .section-divider {
        height: 1px;
        background: linear-gradient(90deg, transparent, #E8EEF2, transparent);
        margin: 2rem 0;
    }

    /* ===== 书籍卡片 - 优化设计 ===== */
    .book-card-container {
        display: flex;
        justify-content: center;
        align-items: stretch;
        gap: 1.5rem;
        margin-bottom: 1.5rem;
    }

    .book-card {
        background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%);
        border-radius: 20px;
        padding: 1.5rem;
        border: 1px solid rgba(102, 126, 234, 0.1);
        transition: all 0.4s cubic-bezier(0.4, 0, 0.2, 1);
        text-align: center;
        flex: 1;
        max-width: 400px;
        display: flex;
        flex-direction: column;
        justify-content: space-between;
        box-shadow: 0 4px 12px rgba(0, 0, 0, 0.05);
        position: relative;
        overflow: hidden;
    }

    .book-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 3px;
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        opacity: 0;
        transition: opacity 0.3s ease;
    }

    .book-card.available {
        cursor: pointer;
    }

    .book-card.available:hover {
        border-color: rgba(102, 126, 234, 0.3);
        box-shadow: 0 8px 24px rgba(102, 126, 234, 0.15);
        transform: translateY(-4px);
    }

    .book-card.available:hover::before {
        opacity: 1;
    }

    .book-card.unavailable {
        opacity: 0.5;
        filter: grayscale(0.3);
    }

    .book-cover {
        font-size: 4.5rem;
        margin: 0.5rem 0;
        opacity: 0.9;
        filter: drop-shadow(0 2px 4px rgba(0, 0, 0, 0.1));
        transition: transform 0.3s ease;
    }

    .book-card.available:hover .book-cover {
        transform: scale(1.1);
    }

    .book-info {
        flex: 1;
        display: flex;
        flex-direction: column;
        justify-content: space-between;
    }

    .book-title {
        font-family: 'Noto Serif SC', serif;
        font-size: 1.15rem;
        font-weight: 700;
        color: #2D3436;
        margin-bottom: 0.5rem;
        line-height: 1.3;
    }

    .book-author {
        font-family: 'Inter', sans-serif;
        font-size: 0.8rem;
        color: #667eea;
        font-weight: 600;
        margin-bottom: 0.75rem;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }

    .book-description {
        font-family: 'Noto Serif SC', serif;
        font-size: 0.85rem;
        color: #636E72;
        line-height: 1.6;
        margin-bottom: 1rem;
        min-height: 2.5rem;
    }

    /* ===== 标签 - 优化设计 ===== */
    .tag-container {
        display: flex;
        gap: 0.4rem;
        flex-wrap: wrap;
        justify-content: center;
        margin-top: 0.75rem;
    }

    .tag {
        display: inline-block;
        padding: 0.3rem 0.75rem;
        background: linear-gradient(145deg, #f0f3f5 0%, #e8eef2 100%);
        color: #2D3436;
        border-radius: 20px;
        font-size: 0.7rem;
        font-weight: 600;
        font-family: 'Inter', sans-serif;
        border: 1px solid rgba(102, 126, 234, 0.1);
        transition: all 0.3s ease;
    }

    .book-card.available:hover .tag {
        background: linear-gradient(145deg, #e8eef2 0%, #dfe6ed 100%);
        border-color: rgba(102, 126, 234, 0.2);
    }

    .tag.highlight {
        background: linear-gradient(145deg, #667eea 0%, #764ba2 100%);
        color: #ffffff;
        border-color: transparent;
    }

    /* ===== 导航按钮 ===== */
    .nav-container {
        display: flex;
        justify-content: center;
        gap: 1rem;
        margin: 2rem 0;
        padding: 1.5rem 0;
        border-top: 1px solid #E8EEF2;
    }

    /* ===== 步骤列表 ===== */
    .step-list {
        list-style: none;
        padding: 0;
        margin: 1.25rem 0;
    }

    .step-item {
        display: flex;
        align-items: flex-start;
        padding: 0.75rem 0;
        border-bottom: 1px solid #E8EEF2;
    }

    .step-item:last-child {
        border-bottom: none;
    }

    .step-number {
        display: flex;
        align-items: center;
        justify-content: center;
        width: 24px;
        height: 24px;
        background: #636E72;
        color: #FFFFFF;
        border-radius: 50%;
        font-weight: 600;
        font-size: 0.85rem;
        margin-right: 0.875rem;
        flex-shrink: 0;
    }

    .step-text {
        font-family: 'Noto Serif SC', serif;
        font-size: 0.95rem;
        line-height: 1.7;
        color: #2D3436;
        flex: 1;
    }

    /* ===== 洞察编号 - 简约线条 ===== */
    .insight-number {
        font-family: 'Inter', sans-serif;
        font-size: 1.8rem;
        font-weight: 600;
        color: #636E72;
        opacity: 0.25;
        margin-bottom: 0.5rem;
        letter-spacing: 0.1em;
    }

    /* ===== 按钮样式 ===== */
    .stButton > button {
        font-family: 'Inter', sans-serif;
        font-weight: 500;
        border-radius: 8px;
        padding: 0.65rem 1.5rem;
        border: none;
        transition: all 0.3s ease;
    }

    .stButton > button[kind="primary"] {
        background: #2D3436;
    }

    /* 收藏按钮样式 - 更优雅的图标按钮 */
    .stButton > button[title*="收藏"],
    .stButton > button[title*="取消"] {
        background: transparent !important;
        border: 1px solid #E8EEF2 !important;
        padding: 0.5rem !important;
        font-size: 1.5rem !important;
        border-radius: 50% !important;
        width: 50px !important;
        height: 50px !important;
        transition: all 0.2s ease !important;
    }

    .stButton > button[title*="收藏"]:hover,
    .stButton > button[title*="取消"]:hover {
        background: #FFF5F5 !important;
        border-color: #FF6B6B !important;
        transform: scale(1.1);
    }

    /* ===== 文本输入框 ===== */
    .stTextArea > div > div > textarea {
        font-family: 'Noto Serif SC', serif;
        font-size: 1rem;
        line-height: 1.8;
        border-radius: 8px;
        border: 1px solid #E8EEF2;
    }

    /* ===== 滚动行为 ===== */
    html {
        scroll-behavior: smooth;
    }

    body {
        scroll-behavior: smooth;
    }

    /* ===== 导出成功提示 ===== */
    .export-success {
        background: #E8F5E9;
        border-left: 4px solid #4CAF50;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
    }

    .export-info {
        background: #E3F2FD;
        border-left: 4px solid #2196F3;
        padding: 1rem;
        border-radius: 8px;
        margin: 1rem 0;
        font-size: 0.9rem;
        line-height: 1.6;
    }

</style>
""", unsafe_allow_html=True)


# ==================== 字体下载函数 ====================

def download_chinese_font():
    """下载中文字体（如果不存在）"""
    font_dir = Path(__file__).parent / 'fonts'
    font_dir.mkdir(exist_ok=True)

    # 使用轻量级的文泉驿字体（约5MB）
    font_file = font_dir / 'wqy-zenhei.ttc'

    if font_file.exists():
        return str(font_file)

    # 尝试下载字体
    try:
        import urllib.request
        import platform

        # 根据平台选择合适的字体
        if platform.system() == 'Windows':
            # Windows 直接使用系统字体
            if Path('C:/Windows/Fonts/msyh.ttc').exists():
                return 'C:/Windows/Fonts/msyh.ttc'
        elif platform.system() == 'Darwin':  # macOS
            if Path('/System/Library/Fonts/PingFang.ttc').exists():
                return '/System/Library/Fonts/PingFang.ttc'

        # Linux 或其他系统，尝试下载
        urls = [
            "https://github.com/googlefonts/noto-cjk/raw/main/Sans/OTF/SimplifiedChinese/NotoSansSC-Regular.otf",
            "https://raw.githubusercontent.com/googlefonts/noto-cjk/main/Sans/OTF/SimplifiedChinese/NotoSansSC-Regular.otf",
        ]

        for url in urls:
            try:
                urllib.request.urlretrieve(url, font_dir / 'NotoSansSC-Regular.otf')
                return str(font_dir / 'NotoSansSC-Regular.otf')
            except:
                continue

    except Exception as e:
        pass

    return None


# ==================== 图片生成函数 ====================

def create_quote_card_image(title, author, quote):
    """生成金句卡片图片 - 优化排版和字体"""
    # 小红书头图尺寸：1080x1440 (3:4比例)
    width_inch = 10.8
    height_inch = 14.4
    dpi = 100

    # 下载或获取中文字体
    chinese_font_path = download_chinese_font()

    # 设置matplotlib的全局字体
    plt.rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial']
    plt.rcParams['axes.unicode_minus'] = False

    chinese_available = False
    if chinese_font_path:
        try:
            font_manager.fontManager.addfont(chinese_font_path)
            font_prop = font_manager.FontProperties(fname=chinese_font_path)
            font_name = font_prop.get_name()
            plt.rcParams['font.sans-serif'] = [font_name, 'DejaVu Sans']
            chinese_available = True
        except:
            chinese_available = False

    # 创建图形
    fig, ax = plt.subplots(figsize=(width_inch, height_inch), dpi=dpi)
    ax.set_xlim(0, 108)
    ax.set_ylim(0, 144)
    ax.axis('off')

    # 绘制白色背景
    ax.add_patch(patches.Rectangle((0, 0), 108, 144, facecolor='white', edgecolor='none'))

    # 绘制顶部紫色纯色条
    ax.add_patch(patches.Rectangle((0, 130), 108, 10, facecolor='#667eea', edgecolor='none'))

    # 绘制金句背景卡片（原来的位置）
    quote_y = 40
    quote_height = 75
    ax.add_patch(patches.FancyBboxPatch((12, quote_y), 84, quote_height,
                                        boxstyle="round,pad=3",
                                        facecolor='#F8F9FA',
                                        edgecolor='#667eea', linewidth=0.4))

    # 绘制装饰线条
    ax.plot([17, 23], [quote_y + 65, quote_y + 65], color='#667eea', linewidth=0.8)
    ax.plot([85, 91], [quote_y + 10, quote_y + 10], color='#667eea', linewidth=0.8)

    # 处理金句文本（分行显示）
    quote_clean = quote.replace('\n', ' ').strip()

    # 简单按字符数分行（每行约12个字符）
    lines = []
    current_line = ""
    for char in quote_clean:
        if len(current_line) < 12:
            current_line += char
        else:
            lines.append(current_line)
            current_line = char
    if current_line:
        lines.append(current_line)

    lines = lines[:4]  # 最多4行

    # 绘制金句文本（垂直居中）
    line_height = 8
    total_height = len(lines) * line_height
    start_y = quote_y + quote_height / 2 + total_height / 2 - 2

    if chinese_available:
        for i, line in enumerate(lines):
            ax.text(54, start_y - i * line_height, line,
                   fontsize=36, color='#2D3436',
                   ha='center', va='center', weight='bold')
    else:
        # 如果中文不可用，显示占位文本
        ax.text(54, start_y, "Deep Reading",
               fontsize=36, color='#2D3436',
               ha='center', va='center', weight='bold')
        ax.text(54, start_y - 8, "Critical Thinking",
               fontsize=28, color='#636E72',
               ha='center', va='center')

    # 绘制标题（在金句框下方）
    if chinese_available:
        ax.text(54, 28, title, fontsize=26, color='#667eea',
                ha='center', va='center', weight='normal')
    else:
        ax.text(54, 28, "QUOTE CARD", fontsize=26, color='#667eea',
                ha='center', va='center', weight='normal')

    # 绘制作者（在标题下方）
    if chinese_available:
        ax.text(54, 23, author, fontsize=16, color='#636E72',
                ha='center', va='center', style='italic')
    else:
        ax.text(54, 23, "By Author", fontsize=16, color='#636E72',
                ha='center', va='center', style='italic')

    # 绘制底部品牌区域
    brand_y = 10
    brand_height = 8

    # 纯色背景矩形
    ax.add_patch(patches.Rectangle((34, brand_y), 40, brand_height,
                                   facecolor='#F8F9FA',
                                   edgecolor='#667eea',
                                   linewidth=0.3))

    # 品牌文本（统一字体大小，避免重叠）
    if chinese_available:
        ax.text(54, brand_y + 5.5, "DeepRead 深读",
               fontsize=22, color='#667eea',
               ha='center', va='center', weight='bold')
        ax.text(54, brand_y + 2.5, "深度阅读 · 沉浸思考",
               fontsize=14, color='#636E72',
               ha='center', va='center')
    else:
        ax.text(54, brand_y + 5.5, "DeepRead",
               fontsize=22, color='#667eea',
               ha='center', va='center', weight='bold')
        ax.text(54, brand_y + 2.5, "Deep Reading",
               fontsize=14, color='#636E72',
               ha='center', va='center')

    # 保存到BytesIO
    buf = BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', pad_inches=0,
                dpi=dpi, facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)

    return buf.getvalue()


def create_reading_poster_image(title, author, emoji, tags, quote, stats):
    """生成阅读海报图片 - 简化版本，避免字体错误"""
    # 下载或获取中文字体
    chinese_font_path = download_chinese_font()

    # 设置matplotlib的全局字体（简单方式）
    plt.rcParams['font.sans-serif'] = ['DejaVu Sans', 'Arial']
    plt.rcParams['axes.unicode_minus'] = False

    chinese_available = False
    if chinese_font_path:
        try:
            # 注册字体
            font_manager.fontManager.addfont(chinese_font_path)
            font_prop = font_manager.FontProperties(fname=chinese_font_path)
            font_name = font_prop.get_name()
            plt.rcParams['font.sans-serif'] = [font_name, 'DejaVu Sans']
            chinese_available = True
        except:
            chinese_available = False

    # 计算高度（减少间距，更紧凑）
    padding = 3
    emoji_h = 6
    title_h = 4
    author_h = 2
    tags_h = 3 if tags else 0
    quote_h = 10
    stats_h = 16
    brand_h = 5

    total_height = padding + emoji_h + title_h + author_h + tags_h + padding + quote_h + padding + stats_h + padding + brand_h

    # 创建图形（宽度60，高度根据内容）
    fig, ax = plt.subplots(figsize=(6, total_height / 10), dpi=100)
    ax.set_xlim(0, 60)
    ax.set_ylim(0, total_height)
    ax.axis('off')

    # 绘制白色背景
    ax.add_patch(patches.Rectangle((0, 0), 60, total_height, facecolor='white', edgecolor='none'))

    y = total_height - padding

    # Emoji
    try:
        ax.text(30, y, emoji, fontsize=28, ha='center', va='top')
    except:
        ax.text(30, y, '📖', fontsize=28, ha='center', va='top')
    y -= emoji_h

    # 标题（缩小字体）
    if chinese_available:
        ax.text(30, y, title, fontsize=22, color='#2D3436',
               ha='center', va='top', weight='bold')
    else:
        ax.text(30, y, "Reading", fontsize=22, color='#2D3436',
               ha='center', va='top', weight='bold')
    y -= title_h

    # 作者（缩小字体）
    if chinese_available:
        ax.text(30, y, author, fontsize=13, color='#636E72',
               ha='center', va='top')
    else:
        ax.text(30, y, "By Author", fontsize=13, color='#636E72',
               ha='center', va='top')
    y -= author_h + 1

    # 标签
    if tags:
        tag_width = 12
        tag_spacing = 1
        total_tags_width = len(tags[:3]) * (tag_width + tag_spacing)

        start_x = max(padding, (60 - total_tags_width) / 2)
        current_x = start_x

        for tag in tags[:3]:
            # 标签背景（使用元组格式）
            ax.add_patch(patches.Rectangle((current_x, y - 3), tag_width, 3.5,
                                          facecolor=(102/255, 126/255, 234/255, 0.1),
                                          edgecolor='#667eea', linewidth=0.15))
            try:
                ax.text(current_x + tag_width / 2, y - 1.5, tag,
                       fontsize=11, color='#667eea',
                       ha='center', va='center')
            except:
                pass
            current_x += tag_width + tag_spacing
        y -= tags_h + 1

    y -= padding

    # 金句区域
    quote_top = y
    quote_bottom = y - quote_h

    ax.add_patch(patches.FancyBboxPatch((padding, quote_bottom), 60 - 2 * padding, quote_h,
                                        boxstyle="round,pad=0.3",
                                        facecolor='#F8F9FA',
                                        edgecolor='#667eea', linewidth=0.3))

    # 金句文本
    quote_clean = quote.replace('\n', ' ').strip()

    # 按字符数分行
    lines = []
    current_line = ""
    for char in quote_clean:
        if len(current_line) < 14:
            current_line += char
        else:
            lines.append(current_line)
            current_line = char
    if current_line:
        lines.append(current_line)

    lines = lines[:3]  # 最多3行

    line_height = 3.5
    total_quote_height = len(lines) * line_height
    quote_start_y = quote_top - (quote_h - total_quote_height) / 2 - 1

    for i, line in enumerate(lines):
        if chinese_available:
            ax.text(30, quote_start_y - i * line_height, line,
                   fontsize=16, color='#2D3436',
                   ha='center', va='top', weight='bold')
        else:
            if i == 0:
                ax.text(30, quote_start_y, "Deep Reading",
                       fontsize=16, color='#2D3436',
                       ha='center', va='top', weight='bold')

    y = quote_bottom - padding

    # 统计区域
    books_read = stats.get('books_read', 0)
    time_text = stats.get('time_display', '0分钟')

    # 已读书籍
    stats_y_start = y
    stats_height = 9

    ax.add_patch(patches.Rectangle((padding, stats_y_start - stats_height), 60 - 2 * padding, stats_height,
                                   facecolor=(102/255, 126/255, 234/255, 0.05),
                                   edgecolor='#667eea', linewidth=0.2))

    # Emoji图标
    ax.text(padding + 2, stats_y_start - 3, '📚', fontsize=18, va='center')

    # 数字
    ax.text(padding + 9, stats_y_start - 3, str(books_read),
           fontsize=32, color='#667eea',
           ha='center', va='center', weight='bold')

    # 标签
    if chinese_available:
        ax.text(padding + 9, stats_y_start - 6, '已读书籍',
               fontsize=12, color='#636E72',
               ha='center', va='center')
    else:
        ax.text(padding + 9, stats_y_start - 6, 'Books Read',
               fontsize=12, color='#636E72',
               ha='center', va='center')

    # 阅读时长
    y = stats_y_start - stats_height - 1.5

    ax.add_patch(patches.Rectangle((padding, y - stats_height), 60 - 2 * padding, stats_height,
                                   facecolor=(118/255, 75/255, 162/255, 0.05),
                                   edgecolor='#764ba2', linewidth=0.2))

    ax.text(padding + 2, y - 3, '⏱️', fontsize=18, va='center')

    try:
        ax.text(padding + 9, y - 3, time_text,
               fontsize=32, color='#764ba2',
               ha='center', va='center', weight='bold')
    except:
        pass

    if chinese_available:
        ax.text(padding + 9, y - 6, '阅读时长',
               fontsize=12, color='#636E72',
               ha='center', va='center')
    else:
        ax.text(padding + 9, y - 6, 'Time Spent',
               fontsize=12, color='#636E72',
               ha='center', va='center')

    # 底部品牌
    brand_y = y - stats_height - 1

    if chinese_available:
        ax.text(30, brand_y, "DeepRead 深读",
               fontsize=14, color='#667eea',
               ha='center', va='center', weight='bold')
        ax.text(30, brand_y - 1.8, "深度阅读 · 沉浸思考",
               fontsize=10, color='#636E72',
               ha='center', va='center')
    else:
        ax.text(30, brand_y, "DeepRead",
               fontsize=14, color='#667eea',
               ha='center', va='center', weight='bold')
        ax.text(30, brand_y - 1.8, "Deep Reading",
               fontsize=10, color='#636E72',
               ha='center', va='center')

    # 保存到BytesIO
    buf = BytesIO()
    plt.savefig(buf, format='png', bbox_inches='tight', pad_inches=0.2,
                dpi=100, facecolor='white', edgecolor='none')
    buf.seek(0)
    plt.close(fig)

    return buf.getvalue()


def generate_quote_card_html(title, author, quote):
    """生成金句卡片的HTML（用于截图备用）"""
    # 处理换行
    quote_display = quote.replace('\n', '<br/>')

    html = f"""
    <div style="
        width: 100%;
        max-width: 540px;
        min-height: 720px;
        margin: 20px auto;
        background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%);
        border-radius: 20px;
        padding: 40px 30px;
        box-shadow: 0 10px 40px rgba(0,0,0,0.1);
        box-sizing: border-box;
        position: relative;
        overflow: hidden;
    ">
        <!-- 顶部装饰条 -->
        <div style="
            position: absolute;
            top: 0;
            left: 0;
            right: 0;
            height: 60px;
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        "></div>

        <!-- 内容区域 -->
        <div style="margin-top: 40px; text-align: center;">

            <!-- 标题 -->
            <h2 style="
                font-family: 'Noto Serif SC', serif;
                font-size: 32px;
                font-weight: 700;
                color: #667eea;
                margin: 0 0 15px 0;
                letter-spacing: 0.02em;
            ">{title}</h2>

            <!-- 作者 -->
            <p style="
                font-family: 'Inter', sans-serif;
                font-size: 18px;
                color: #636E72;
                margin: 0 0 40px 0;
            ">{author}</p>

            <!-- 金句卡片 -->
            <div style="
                background: #F8F9FA;
                border-radius: 16px;
                padding: 40px 30px;
                margin: 30px 0;
                border-left: 4px solid #667eea;
                position: relative;
            ">
                <!-- 装饰线条 -->
                <div style="
                    position: absolute;
                    top: 20px;
                    left: 20px;
                    width: 40px;
                    height: 3px;
                    background: #667eea;
                "></div>
                <div style="
                    position: absolute;
                    bottom: 20px;
                    right: 20px;
                    width: 40px;
                    height: 3px;
                    background: #667eea;
                "></div>

                <!-- 金句文本 -->
                <p style="
                    font-family: 'Noto Serif SC', serif;
                    font-size: 26px;
                    font-weight: 600;
                    color: #2D3436;
                    line-height: 1.8;
                    margin: 0;
                ">{quote_display}</p>
            </div>

            <!-- 品牌区域 -->
            <div style="
                margin-top: 60px;
                padding-top: 30px;
                border-top: 2px solid #E8EEF2;
            ">
                <div style="
                    display: inline-block;
                    padding: 15px 30px;
                    background: linear-gradient(135deg, rgba(102, 126, 234, 0.1) 0%, rgba(118, 75, 162, 0.1) 100%);
                    border-radius: 30px;
                ">
                    <p style="
                        font-family: 'Noto Serif SC', serif;
                        font-size: 20px;
                        font-weight: 700;
                        color: #667eea;
                        margin: 0 0 5px 0;
                    ">DeepRead 深读</p>
                    <p style="
                        font-family: 'Inter', sans-serif;
                        font-size: 14px;
                        color: #636E72;
                        margin: 0;
                    ">深度阅读 · 沉浸思考</p>
                </div>
            </div>
        </div>
    </div>
    """
    return html


def init_session_state():
    """初始化session state"""
    if "current_book" not in st.session_state:
        st.session_state.current_book = None
    if "current_content" not in st.session_state:
        st.session_state.current_content = None
    if "current_section" not in st.session_state:
        st.session_state.current_section = "library"
    if "notes" not in st.session_state:
        st.session_state.notes = {}
    if "page_rerun" not in st.session_state:
        st.session_state.page_rerun = 0

    # 用户状态管理（新增）
    if "is_first_visit" not in st.session_state:
        st.session_state.is_first_visit = True
    if "trial_start_date" not in st.session_state:
        st.session_state.trial_start_date = datetime.now().date()
    if "user_tier" not in st.session_state:
        st.session_state.user_tier = "trial"  # trial, free, premium
    if "guide_completed" not in st.session_state:
        st.session_state.guide_completed = False
    if "guide_step" not in st.session_state:
        st.session_state.guide_step = 0  # 当前引导步骤（0-3）

    # 阅读统计数据
    if "reading_stats" not in st.session_state:
        st.session_state.reading_stats = {
            "total_books_read": set(),  # 已读过的书名集合
            "total_reading_time": 0,     # 总阅读时长（秒）
            "last_read_time": None       # 最后阅读时间
        }

    # 实践任务追踪数据
    if "practice_tracker" not in st.session_state:
        st.session_state.practice_tracker = {}  # 格式: {book_title: {week: {day: completed}}}

    # 成就系统数据
    if "achievements" not in st.session_state:
        st.session_state.achievements = {
            "unlocked": [],          # 已解锁的成就ID列表
            "notifications": [],     # 成就通知队列
            "last_check_time": None  # 上次检查成就的时间
        }

    # 阅读进度追踪
    if "reading_progress" not in st.session_state:
        st.session_state.reading_progress = {}  # 格式: {book_id: {"current_chapter": int, "progress_percent": float, "last_read": datetime}}

    # 用户账户系统（P2功能）
    if "user_account" not in st.session_state:
        st.session_state.user_account = {
            "logged_in": False,          # 是否登录
            "user_id": None,             # 用户ID
            "username": None,            # 用户名
            "email": None,               # 邮箱
            "created_at": None,          # 注册时间
            "last_sync": None,           # 最后同步时间
            "preferences": {             # 用户偏好
                "theme": "light",        # 主题：light/dark
                "notification_enabled": True,  # 通知开关
                "reminder_time": "09:00",    # 提醒时间
                "export_format": "markdown"   # 默认导出格式
            }
        }

    # 智能复习提醒（P2功能）
    if "review_reminders" not in st.session_state:
        st.session_state.review_reminders = {
            "enabled": False,            # 是否启用复习提醒
            "books_to_review": [],       # 需要复习的书籍列表
            "review_schedule": {},       # 复习计划: {book_id: review_dates}
            "last_review_check": None    # 上次检查复习的时间
        }


# ==================== 用户管理相关函数 ====================

def get_trial_days_remaining():
    """计算试用剩余天数"""
    if st.session_state.user_tier != "trial":
        return 0

    trial_start = st.session_state.trial_start_date
    days_passed = (datetime.now().date() - trial_start).days
    remaining = 7 - days_passed
    return max(0, remaining)


def show_trial_notice():
    """显示试用提醒横幅（如果需要）"""
    if st.session_state.user_tier == "trial":
        days_remaining = get_trial_days_remaining()

        if days_remaining > 0:
            st.info(f"🎁 7天深度版免费试用中，还剩 {days_remaining} 天")
        elif days_remaining == 0:
            st.warning("⏰ 试用已到期，升级深度版解锁更多功能")
            st.markdown("""
            <div style="text-align: center; margin: 1rem 0;">
                <a href="#upgrade" style="color: #667eea; text-decoration: none; font-weight: 600;">
                    🔓 查看深度版功能对比 →
                </a>
            </div>
            """, unsafe_allow_html=True)


# ==================== 成就系统 ====================

# 成就定义配置
ACHIEVEMENTS_DEFINITIONS = {
    "first_book": {
        "id": "first_book",
        "name": "阅读萌芽",
        "icon": "🌱",
        "description": "完成第1本书",
        "condition": lambda stats: len(stats["total_books_read"]) >= 1,
        "tier": 1
    },
    "five_books": {
        "id": "five_books",
        "name": "阅读爱好者",
        "icon": "📚",
        "description": "完成5本书",
        "condition": lambda stats: len(stats["total_books_read"]) >= 5,
        "tier": 2
    },
    "ten_books": {
        "id": "ten_books",
        "name": "深度阅读者",
        "icon": "👑",
        "description": "完成10本书",
        "condition": lambda stats: len(stats["total_books_read"]) >= 10,
        "tier": 3
    },
    "first_note": {
        "id": "first_note",
        "name": "动笔思考",
        "icon": "✍️",
        "description": "记录第一条实践笔记",
        "condition": lambda stats, notes: sum(len(practices) for practices in notes.values()) >= 1,
        "tier": 1
    },
    "ten_notes": {
        "id": "ten_notes",
        "name": "思考积累",
        "icon": "💡",
        "description": "记录10条实践笔记",
        "condition": lambda stats, notes: sum(len(practices) for practices in notes.values()) >= 10,
        "tier": 2
    },
    "first_hour": {
        "id": "first_hour",
        "name": "入门时光",
        "icon": "⏱️",
        "description": "累计阅读1小时",
        "condition": lambda stats: stats["total_reading_time"] >= 3600,
        "tier": 1
    },
    "ten_hours": {
        "id": "ten_hours",
        "name": "投入阅读",
        "icon": "⌛",
        "description": "累计阅读10小时",
        "condition": lambda stats: stats["total_reading_time"] >= 36000,
        "tier": 2
    },
    "streak_3_days": {
        "id": "streak_3_days",
        "name": "连续阅读",
        "icon": "🔥",
        "description": "连续3天阅读",
        "condition": lambda stats: len(stats.get("daily_progress", {})) >= 3,
        "tier": 2
    },
    "first_reflection": {
        "id": "first_reflection",
        "name": "深度反思",
        "icon": "🤔",
        "description": "写下第一篇深度反思",
        "condition": lambda stats, reflections: sum(len(refs) for refs in reflections.values()) >= 1,
        "tier": 1
    }
}


def check_and_unlock_achievements():
    """检查并解锁成就"""
    stats = st.session_state.reading_stats
    notes = st.session_state.practices
    reflections = st.session_state.reflections

    newly_unlocked = []

    for achievement_id, achievement in ACHIEVEMENTS_DEFINITIONS.items():
        # 跳过已解锁的成就
        if achievement_id in st.session_state.achievements["unlocked"]:
            continue

        # 检查解锁条件
        try:
            # 根据成就类型调用不同的条件函数
            if achievement_id in ["first_note", "ten_notes"]:
                is_unlocked = achievement["condition"](stats, notes)
            elif achievement_id == "first_reflection":
                is_unlocked = achievement["condition"](stats, reflections)
            else:
                is_unlocked = achievement["condition"](stats)

            if is_unlocked:
                # 解锁成就
                st.session_state.achievements["unlocked"].append(achievement_id)
                newly_unlocked.append(achievement)

                # 添加到通知队列
                st.session_state.achievements["notifications"].append({
                    "achievement_id": achievement_id,
                    "timestamp": datetime.now(),
                    "shown": False
                })
        except Exception as e:
            # 静默失败，避免影响用户体验
            pass

    return newly_unlocked


def show_achievement_notifications():
    """显示成就解锁通知"""
    notifications = st.session_state.achievements["notifications"]

    for i, notification in enumerate(notifications):
        if not notification["shown"]:
            achievement_id = notification["achievement_id"]
            achievement = ACHIEVEMENTS_DEFINITIONS.get(achievement_id)

            if achievement:
                # 显示成就解锁通知
                st.markdown(f"""
                <div style="
                    position: fixed;
                    top: 20px;
                    right: 20px;
                    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    color: white;
                    padding: 1.5rem 2rem;
                    border-radius: 12px;
                    box-shadow: 0 10px 30px rgba(102, 126, 234, 0.3);
                    z-index: 9999;
                    animation: slideIn 0.5s ease-out;
                    min-width: 300px;
                ">
                    <div style="font-size: 0.8rem; opacity: 0.9; margin-bottom: 0.5rem;">
                        🎉 成就解锁
                    </div>
                    <div style="display: flex; align-items: center; gap: 1rem;">
                        <div style="font-size: 2.5rem;">{achievement['icon']}</div>
                        <div>
                            <div style="font-size: 1.1rem; font-weight: 700; margin-bottom: 0.25rem;">
                                {achievement['name']}
                            </div>
                            <div style="font-size: 0.8rem; opacity: 0.9;">
                                {achievement['description']}
                            </div>
                        </div>
                    </div>
                </div>

                <style>
                @keyframes slideIn {{
                    from {{
                        transform: translateX(400px);
                        opacity: 0;
                    }}
                    to {{
                        transform: translateX(0);
                        opacity: 1;
                    }}
                }}
                </style>
                """, unsafe_allow_html=True)

                # 标记为已显示
                st.session_state.achievements["notifications"][i]["shown"] = True

                # 自动关闭通知（3秒后）
                import time
                time.sleep(3)


def get_achievement_progress():
    """获取成就进度信息"""
    stats = st.session_state.reading_stats
    notes = st.session_state.practices
    reflections = st.session_state.reflections

    progress_info = {}

    for achievement_id, achievement in ACHIEVEMENTS_DEFINITIONS.items():
        is_unlocked = achievement_id in st.session_state.achievements["unlocked"]

        # 计算进度
        try:
            if achievement_id == "first_book":
                current = len(stats["total_books_read"])
                target = 1
            elif achievement_id == "five_books":
                current = len(stats["total_books_read"])
                target = 5
            elif achievement_id == "ten_books":
                current = len(stats["total_books_read"])
                target = 10
            elif achievement_id == "first_note":
                current = sum(len(practices) for practices in notes.values())
                target = 1
            elif achievement_id == "ten_notes":
                current = sum(len(practices) for practices in notes.values())
                target = 10
            elif achievement_id == "first_hour":
                current = stats["total_reading_time"] // 3600
                target = 1
            elif achievement_id == "ten_hours":
                current = stats["total_reading_time"] // 3600
                target = 10
            elif achievement_id == "streak_3_days":
                current = len(stats.get("daily_progress", {}))
                target = 3
            elif achievement_id == "first_reflection":
                current = sum(len(refs) for refs in reflections.values())
                target = 1
            else:
                current = 0
                target = 1

            progress_info[achievement_id] = {
                "unlocked": is_unlocked,
                "current": min(current, target),
                "target": target,
                "percent": min(int((current / target) * 100) if target > 0 else 0, 100)
            }
        except Exception:
            progress_info[achievement_id] = {
                "unlocked": is_unlocked,
                "current": 0,
                "target": 1,
                "percent": 0
            }

    return progress_info


def update_reading_progress(book_id, chapter_index, total_chapters):
    """更新阅读进度"""
    if book_id not in st.session_state.reading_progress:
        st.session_state.reading_progress[book_id] = {}

    progress_percent = int((chapter_index / total_chapters) * 100) if total_chapters > 0 else 0

    st.session_state.reading_progress[book_id] = {
        "current_chapter": chapter_index,
        "total_chapters": total_chapters,
        "progress_percent": progress_percent,
        "last_read": datetime.now()
    }


# ==================== 智能复习提醒系统 ====================

def schedule_review(book_id, book_title):
    """为已完成的书籍安排复习计划（基于艾宾浩斯遗忘曲线）"""
    if not st.session_state.review_reminders["enabled"]:
        return

    # 艾宾浩斯遗忘曲线复习点：1天、3天、7天、15天、30天
    review_intervals = [1, 3, 7, 15, 30]

    review_dates = []
    completion_date = datetime.now().date()

    for interval in review_intervals:
        review_date = completion_date + timedelta(days=interval)
        review_dates.append({
            "interval": interval,
            "date": review_date.strftime("%Y-%m-%d"),
            "completed": False
        })

    st.session_state.review_reminders["review_schedule"][book_id] = {
        "title": book_title,
        "completion_date": completion_date.strftime("%Y-%m-%d"),
        "reviews": review_dates
    }

    # 添加到待复习列表
    if book_id not in st.session_state.review_reminders["books_to_review"]:
        st.session_state.review_reminders["books_to_review"].append(book_id)


def check_review_reminders():
    """检查是否有需要复习的书籍"""
    if not st.session_state.review_reminders["enabled"]:
        return []

    today = datetime.now().date()
    due_reviews = []

    for book_id in st.session_state.review_reminders["books_to_review"]:
        schedule = st.session_state.review_reminders["review_schedule"].get(book_id)
        if not schedule:
            continue

        for review in schedule["reviews"]:
            # 检查是否到期且未完成
            review_date = datetime.strptime(review["date"], "%Y-%m-%d").date()
            if review_date <= today and not review["completed"]:
                due_reviews.append({
                    "book_id": book_id,
                    "title": schedule["title"],
                    "interval": review["interval"],
                    "date": review["date"]
                })

    return due_reviews


def mark_review_complete(book_id, interval_days):
    """标记某次复习已完成"""
    schedule = st.session_state.review_reminders["review_schedule"].get(book_id)
    if not schedule:
        return

    for review in schedule["reviews"]:
        if review["interval"] == interval_days:
            review["completed"] = True
            break

    # 检查是否所有复习都已完成
    all_completed = all(review["completed"] for review in schedule["reviews"])
    if all_completed:
        # 从待复习列表中移除
        if book_id in st.session_state.review_reminders["books_to_review"]:
            st.session_state.review_reminders["books_to_review"].remove(book_id)


def show_review_reminder_panel():
    """显示复习提醒面板"""
    if not st.session_state.review_reminders["enabled"]:
        return

    due_reviews = check_review_reminders()

    if not due_reviews:
        return

    st.markdown("""
    <div style="background: linear-gradient(135deg, #ffeaa7 0%, #fdcb6e 100%);
                padding: 1.5rem; border-radius: 12px; margin: 1rem 0;
                border-left: 4px solid #f39c12;">
        <div style="font-size: 1rem; font-weight: 600; color: #2D3436; margin-bottom: 0.5rem;">
            📚 复习提醒
        </div>
    </div>
    """, unsafe_allow_html=True)

    for review in due_reviews:
        interval_text = {
            1: "第1次复习（1天后）",
            3: "第2次复习（3天后）",
            7: "第3次复习（1周后）",
            15: "第4次复习（2周后）",
            30: "第5次复习（1个月后）"
        }.get(review["interval"], f"{review['interval']}天后")

        st.markdown(f"""
        <div style="background: #FFF3CD; padding: 1rem; border-radius: 8px; margin-bottom: 0.5rem;">
            <div style="font-weight: 600; margin-bottom: 0.25rem;">📖 {review['title']}</div>
            <div style="font-size: 0.85rem; color: #636E72;">{interval_text} - 到期日: {review['date']}</div>
        </div>
        """, unsafe_allow_html=True)


def update_reading_progress(book_id, chapter_index, total_chapters):
    """更新阅读进度"""
    if book_id not in st.session_state.reading_progress:
        st.session_state.reading_progress[book_id] = {}

    progress_percent = int((chapter_index / total_chapters) * 100) if total_chapters > 0 else 0

    st.session_state.reading_progress[book_id] = {
        "current_chapter": chapter_index,
        "total_chapters": total_chapters,
        "progress_percent": progress_percent,
        "last_read": datetime.now()
    }


def show_welcome_page():
    """显示首次访问欢迎页"""
    # 使用更稳定的单行HTML格式
    html_content = '<div style="text-align: center; padding: 4rem 2rem; background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); border-radius: 20px; margin: 2rem 0; color: white;"><h1 style="font-family: \'Noto Serif SC\', serif; font-size: 3rem; font-weight: 700; margin: 0 0 1rem 0; color: white;">开始你的深度阅读之旅 🧠</h1><p style="font-size: 1.2rem; margin: 0 0 2rem 0; opacity: 0.95;">不只是阅读，更是深度思考和行动</p><div style="display: flex; justify-content: center; gap: 2rem; flex-wrap: wrap; margin: 2rem 0;"><div style="flex: 1; min-width: 200px; padding: 1.5rem;"><div style="font-size: 3rem; margin-bottom: 0.5rem;">📖</div><div style="font-size: 1.1rem; font-weight: 600;">精选书籍</div><div style="font-size: 0.9rem; opacity: 0.85;">个人成长 · 认知提升</div></div><div style="flex: 1; min-width: 200px; padding: 1.5rem;"><div style="font-size: 3rem; margin-bottom: 0.5rem;">🎯</div><div style="font-size: 1.1rem; font-weight: 600;">实践追踪</div><div style="font-size: 0.9rem; opacity: 0.85;">30天习惯养成</div></div><div style="flex: 1; min-width: 200px; padding: 1.5rem;"><div style="font-size: 3rem; margin-bottom: 0.5rem;">💡</div><div style="font-size: 1.1rem; font-weight: 600;">深度思考</div><div style="font-size: 0.9rem; opacity: 0.85;">反思与输出</div></div></div><div style="background: rgba(255, 255, 255, 0.15); padding: 1rem 2rem; border-radius: 12px; margin: 2rem 0;"><div style="font-size: 1.2rem; margin-bottom: 0.5rem;">🎁 7天深度版免费试用</div><div style="font-size: 0.95rem; opacity: 0.9;">云同步 · 数据统计 · 智能推荐</div></div></div>'

    st.markdown(html_content, unsafe_allow_html=True)

    if st.button("开始探索 🚀", use_container_width=True, key="start_exploring"):
        st.session_state.is_first_visit = False
        st.session_state.guide_step = 1
        st.rerun()


def show_guide_bubble():
    """显示新手引导气泡"""
    if not st.session_state.is_first_visit and st.session_state.guide_step < 4:

        guide_steps = [
            {
                "step": 1,
                "position": "书库页面",
                "message": "👆 选择一本书开始阅读，点击卡片进入导读页",
                "target": "书籍卡片"
            },
            {
                "step": 2,
                "position": "导读页",
                "message": "💡 阅读完内容后，不要忘记记录实践计划和反思思考",
                "target": "导读页"
            },
            {
                "step": 3,
                "position": "侧边栏",
                "message": "📊 点击侧边栏查看你的阅读统计和成就",
                "target": "侧边栏"
            }
        ]

        current_step = st.session_state.guide_step
        if current_step < len(guide_steps):
            step_info = guide_steps[current_step - 1]

            # 使用更简洁的单行HTML格式
            bubble_html = f"""
            <div style="position: fixed; bottom: 20px; right: 20px; z-index: 9999; background: white; padding: 1rem 1.5rem; border-radius: 12px; box-shadow: 0 4px 20px rgba(0,0,0,0.15); border-left: 4px solid #667eea; max-width: 300px;">
                <div style="font-size: 1rem; font-weight: 600; color: #2D3436; margin-bottom: 0.5rem;">
                    {step_info['message']}
                </div>
                <div style="font-size: 0.85rem; color: #636E72; margin-bottom: 0.75rem;">
                    目标：{step_info['target']}
                </div>
            </div>
            """
            st.markdown(bubble_html, unsafe_allow_html=True)

            # 使用Streamlit原生按钮而不是HTML button
            if st.button(f"知道了（{current_step}/3）", key=f"guide_step_{current_step}"):
                st.session_state.guide_step += 1
                st.rerun()


# ==================== 书籍数据 ====================
BOOKS_DATA = [
    {
        "title": "原子习惯",
        "author": "詹姆斯·克利尔",
        "description": "微小改变如何通过时间的复利，带来人生的巨大转变。每天进步1%，一年后你会进步37倍。",
        "tags": ["个人成长"],
        "available": True,
        "emoji": "📖"
    },
    {
        "title": "思考，快与慢",
        "author": "丹尼尔·卡尼曼",
        "description": "理解人类思维的双系统，认识直觉与理性的真相。系统1快速直觉，系统2缓慢理性。",
        "tags": ["认知提升"],
        "available": True,
        "emoji": "🧠"
    },
    {
        "title": "刻意练习",
        "author": "安德斯·艾利克森",
        "description": "揭秘天才的秘密：卓越不是天赋，而是正确的练习方法。突破1万小时的误区，用刻意练习在任何领域达到顶尖水平。",
        "tags": ["个人成长"],
        "available": True,
        "emoji": "🎯"
    },
    {
        "title": "深度工作",
        "author": "卡尔·纽波特",
        "description": "在分心的世界中，深度专注工作的能力是日益稀缺的资产。掌握深度工作法则，在浮躁的世界创造真正的价值。",
        "tags": ["职场发展"],
        "available": True,
        "emoji": "💼"
    },
    {
        "title": "原则",
        "author": "雷·达里欧",
        "description": "世界最大对冲基金创始人的人生智慧。如何通过原则化思维，在工作和生活中做出更好的决策，实现个人与组织的进化。",
        "tags": ["职场发展"],
        "available": True,
        "emoji": "📐"
    },
    {
        "title": "心流",
        "author": "米哈里·契克森米哈赖",
        "description": "最优体验的秘密。为什么有些人能在艰难困苦中找到乐趣，而有些人在优越环境中却感到空虚？幸福不是外在条件，而是内在秩序的建立。",
        "tags": ["个人成长"],
        "available": True,
        "emoji": "🌊"
    },
    {
        "title": "影响力",
        "author": "罗伯特·西奥迪尼",
        "description": "说服的心理学。为什么有些人能轻松影响他人，而你却总是被说服？掌握互惠、稀缺、社会认同等6大说服原则，既保护自己不被套路，也能道德地影响他人。",
        "tags": ["认知提升"],
        "available": True,
        "emoji": "🎭"
    },
    {
        "title": "终身成长",
        "author": "卡罗尔·德韦克",
        "description": "思维模式的力量。为什么有些人遇到挫折就放弃，而有些人越挫越勇？固定型思维vs成长型思维，这个小小的信念差异，决定了你的一生。",
        "tags": ["个人成长"],
        "available": True,
        "emoji": "🌱"
    }
]


def scroll_to_top():
    """滚动到页面顶部"""
    st.markdown("""
<script>
    // 多种方法确保滚动到顶部
    function scrollToTop() {
        // 方法1：window.scrollTo
        window.scrollTo({ top: 0, behavior: 'smooth' });

        // 方法2：document.documentElement
        document.documentElement.scrollTop = 0;

        // 方法3：document.body
        document.body.scrollTop = 0;

        // 方法4：查找主容器
        const mainElement = document.querySelector('.main');
        if (mainElement) {
            mainElement.scrollTop = 0;
        }
    }

    // 页面加载后立即执行
    scrollToTop();

    // 延迟再次执行（确保渲染完成）
    setTimeout(scrollToTop, 100);
    setTimeout(scrollToTop, 300);
</script>
""", unsafe_allow_html=True)


def render_progress_bar(current_section):
    """
    显示阅读进度条

    current_section: 当前所在部分
    - 'intro': 引言
    - 'insights': 洞察
    - 'practice': 实践
    - 'reflection': 反思
    """
    sections = {
        'intro': {'name': '引言', 'icon': '📖', 'progress': 20},
        'insights': {'name': '洞察', 'icon': '💡', 'progress': 40},
        'practice': {'name': '实践', 'icon': '✅', 'progress': 60},
        'reflection': {'name': '反思', 'icon': '🤔', 'progress': 80}
    }

    # 获取当前部分的进度
    current = sections.get(current_section, sections['intro'])

    # 创建进度条HTML
    progress_bar = f"""
<div style="background: #F7F9FC; padding: 0.8rem 1rem; border-radius: 8px; margin-bottom: 1.5rem;">
    <div style="display: flex; align-items: center; justify-content: space-between; margin-bottom: 0.5rem;">
        <div style="font-weight: 600; color: #1F2937; font-size: 0.9rem;">阅读进度</div>
        <div style="color: #6B7280; font-size: 0.8rem;">{current['progress']}%</div>
    </div>
    <div style="background: #E5E7EB; height: 8px; border-radius: 4px; overflow: hidden;">
        <div style="background: linear-gradient(90deg, #667eea 0%, #764ba2 100%); height: 100%; width: {current['progress']}%; transition: width 0.3s ease;"></div>
    </div>
    <div style="display: flex; justify-content: space-between; margin-top: 0.5rem;">
        <div style="font-size: 0.75rem; color: {'#667eea; font-weight: 600;' if current_section == 'intro' else '#9CA3AF'}">📖 引言</div>
        <div style="font-size: 0.75rem; color: {'#667eea; font-weight: 600;' if current_section == 'insights' else '#9CA3AF'}">💡 洞察</div>
        <div style="font-size: 0.75rem; color: {'#667eea; font-weight: 600;' if current_section == 'practice' else '#9CA3AF'}">✅ 实践</div>
        <div style="font-size: 0.75rem; color: {'#667eea; font-weight: 600;' if current_section == 'reflection' else '#9CA3AF'}">🤔 反思</div>
    </div>
</div>
"""

    st.markdown(progress_bar, unsafe_allow_html=True)


def generate_markdown(content, notes):
    """生成Markdown格式的学习笔记"""
    book_title = content["title"]
    author = content["author"]

    # 获取当前日期
    today = datetime.now().strftime("%Y年%m月%d日")

    md = f"""# {book_title} - 学习笔记

**作者**: {author}
**阅读日期**: {today}
**来源**: DeepRead 深度阅读

---

## 📖 引言

### {content["introduction"]["title"]}

{content["introduction"]["subtitle"]}

#### 为什么要读这本书

{content["introduction"]["why_read"]}

---

## 💡 核心洞察

### {content["core_thinking"]["title"]}

{content["core_thinking"]["subtitle"]}

"""

    # 添加每个洞察
    for idx, insight in enumerate(content["core_thinking"]["insights"], 1):
        md += f"""#### {idx}. {insight["title"]}

**核心观点**:
{insight["core_idea"]}

**为什么这很重要**:
{insight["why_matters"]}

"""
        if insight.get("example"):
            md += f"""**现实中的样子**:
{insight["example"]}

"""
        if insight.get("question"):
            md += f"""**想一想**: {insight["question"]}

"""
        md += "---\n\n"

    # 添加实践步骤
    md += f"""## ✅ 实践步骤

### {content["practice"]["title"]}

{content["practice"]["subtitle"]}

"""

    for action in content["practice"]["actions"]:
        md += f"""#### {action["title"]}

{action["description"]}

**步骤**:
"""
        for step in action.get("steps", []):
            md += f"- {step}\n"
        md += "\n"

    # 添加反思和用户笔记
    md += f"""## 🤔 反思与思考

### {content["reflection"]["title"]}

{content["reflection"]["subtitle"]}

"""

    for idx, question in enumerate(content["reflection"]["questions"], 1):
        md += f"""#### 问题 {idx}

{question["text"]}

**提示**: {question["hint"]}

"""
        # 添加用户的笔记
        note_key = f"q{idx}"
        if notes.get(note_key):
            md += f"""**我的思考**:

{notes[note_key]}

"""
        else:
            md += "**我的思考**: *（暂未填写）*\n\n"
        md += "---\n\n"

    # 添加金句
    md += """## 📝 值得记住的话

"""
    for quote in content["quotes"]:
        md += f"""> {quote}

"""

    md += f"""
---

*本笔记由 [DeepRead](https://github.com) 深度阅读工具生成*
"""

    return md


def generate_notes_only(content, notes):
    """只生成用户笔记的Markdown"""
    book_title = content["title"]
    author = content["author"]
    today = datetime.now().strftime("%Y年%m月%d日")

    md = f"""# {book_title} - 我的阅读笔记

**作者**: {author}
**记录日期**: {today}
**来源**: DeepRead 深度阅读

---

## 🤔 我的思考与反思

"""

    has_notes = False
    for idx, question in enumerate(content["reflection"]["questions"], 1):
        note_key = f"q{idx}"
        if notes.get(note_key):
            has_notes = True
            md += f"""### 问题 {idx}

{question["text"]}

**我的答案**:

{notes[note_key]}

---

"""

    if not has_notes:
        md += "*还没有填写任何笔记。回到阅读页面填写问题后，即可导出笔记。*\n\n"

    return md


# ==================== 高级导出功能 ====================

def generate_pdf_bytes(content, notes, include_full_content=True):
    """生成PDF格式的学习笔记"""
    if not PDF_SUPPORT:
        return None

    buffer = BytesIO()

    # 创建PDF文档
    doc = SimpleDocTemplate(buffer, pagesize=A4,
                            rightMargin=2*cm, leftMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    # 样式
    styles = getSampleStyleSheet()

    # 自定义样式
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#2D3436'),
        spaceAfter=30,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )

    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=16,
        textColor=colors.HexColor('#667eea'),
        spaceAfter=12,
        spaceBefore=20,
        fontName='Helvetica-Bold'
    )

    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        textColor=colors.HexColor('#2D3436'),
        spaceAfter=10,
        leading=16,
        fontName='Helvetica'
    )

    # 构建内容
    story = []

    # 标题
    book_title = content["title"]
    author = content["author"]
    today = datetime.now().strftime("%Y年%m月%d日")

    story.append(Paragraph(f"{book_title}", title_style))
    story.append(Spacer(1, 0.3*cm))
    story.append(Paragraph(f"<b>作者</b>: {author}", normal_style))
    story.append(Paragraph(f"<b>阅读日期</b>: {today}", normal_style))
    story.append(Paragraph(f"<b>来源</b>: DeepRead 深度阅读", normal_style))
    story.append(Spacer(1, 0.5*cm))

    # 分隔线
    story.append(Spacer(1, 0.3*cm))

    # 如果包含完整内容
    if include_full_content:
        # 引言部分
        story.append(Paragraph("📖 引言", heading_style))
        intro = content.get("introduction", {})
        if intro:
            story.append(Paragraph(f"<b>{intro.get('title', '')}</b>", normal_style))
            story.append(Paragraph(intro.get('subtitle', ''), normal_style))
            story.append(Spacer(1, 0.3*cm))

            if intro.get("content"):
                for para in intro["content"]:
                    story.append(Paragraph(para, normal_style))
                    story.append(Spacer(1, 0.2*cm))

        story.append(Spacer(1, 0.5*cm))

        # 核心洞察
        story.append(Paragraph("💡 核心洞察", heading_style))
        insights = content.get("insights", {})
        if insights:
            story.append(Paragraph(f"<b>{insights.get('title', '')}</b>", normal_style))
            story.append(Paragraph(insights.get('subtitle', ''), normal_style))
            story.append(Spacer(1, 0.3*cm))

            if insights.get("key_points"):
                for idx, point in enumerate(insights["key_points"], 1):
                    story.append(Paragraph(f"{idx}. {point.get('title', '')}", normal_style))
                    if point.get("description"):
                        story.append(Paragraph(point["description"], normal_style))
                    story.append(Spacer(1, 0.2*cm))

            story.append(Spacer(1, 0.3*cm))

            if insights.get("framework"):
                story.append(Paragraph("<b>核心框架</b>", normal_style))
                for item in insights["framework"]:
                    story.append(Paragraph(f"• {item}", normal_style))

        story.append(Spacer(1, 0.5*cm))

        # 实践行动
        story.append(Paragraph("✍️ 实践行动", heading_style))
        practice = content.get("practice", {})
        if practice:
            story.append(Paragraph(f"<b>{practice.get('title', '')}</b>", normal_style))
            story.append(Paragraph(practice.get('subtitle', ''), normal_style))
            story.append(Spacer(1, 0.3*cm))

            if practice.get("actions"):
                for idx, action in enumerate(practice["actions"], 1):
                    story.append(Paragraph(f"<b>步骤 {idx}</b>: {action.get('title', '')}", normal_style))
                    if action.get("description"):
                        story.append(Paragraph(action["description"], normal_style))
                    if action.get("steps"):
                        for step in action["steps"]:
                            story.append(Paragraph(f"  • {step}", normal_style))
                    story.append(Spacer(1, 0.2*cm))

        story.append(PageBreak())

    # 我的思考与反思
    story.append(Paragraph("🤔 我的思考与反思", heading_style))
    story.append(Spacer(1, 0.3*cm))

    has_notes = False
    for idx, question in enumerate(content.get("reflection", {}).get("questions", []), 1):
        note_key = f"q{idx}"
        if notes.get(note_key):
            has_notes = True
            story.append(Paragraph(f"<b>问题 {idx}</b>", normal_style))
            story.append(Paragraph(question.get("text", ""), normal_style))
            story.append(Spacer(1, 0.2*cm))
            story.append(Paragraph("<b>我的答案</b>:", normal_style))
            story.append(Paragraph(notes[note_key], normal_style))
            story.append(Spacer(1, 0.5*cm))

    if not has_notes:
        story.append(Paragraph("<i>还没有填写任何笔记。</i>", normal_style))

    # 金句摘录
    if content.get("quotes"):
        story.append(PageBreak())
        story.append(Paragraph("💎 值得记住的话", heading_style))
        story.append(Spacer(1, 0.3*cm))
        for quote in content["quotes"]:
            story.append(Paragraph(f"<i>{quote}</i>", normal_style))
            story.append(Spacer(1, 0.3*cm))

    # 生成PDF
    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()

    return pdf_bytes


def generate_word_bytes(content, notes, include_full_content=True):
    """生成Word格式的学习笔记"""
    if not WORD_SUPPORT:
        return None

    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # 标题
    title = doc.add_heading(content["title"], 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元信息
    p = doc.add_paragraph()
    p.add_run(f"作者: {content['author']}\n")
    p.add_run(f"阅读日期: {datetime.now().strftime('%Y年%m月%d日')}\n")
    p.add_run("来源: DeepRead 深度阅读")
    p.style = 'Normal'

    doc.add_paragraph('─' * 50)

    # 完整内容
    if include_full_content:
        # 引言
        doc.add_heading('📖 引言', 1)
        intro = content.get("introduction", {})
        if intro:
            doc.add_heading(intro.get('title', ''), 2)
            doc.add_paragraph(intro.get('subtitle', ''))

            if intro.get("content"):
                for para in intro["content"]:
                    doc.add_paragraph(para)

        # 核心洞察
        doc.add_heading('💡 核心洞察', 1)
        insights = content.get("insights", {})
        if insights:
            doc.add_heading(insights.get('title', ''), 2)
            doc.add_paragraph(insights.get('subtitle', ''))

            if insights.get("key_points"):
                for idx, point in enumerate(insights["key_points"], 1):
                    p = doc.add_paragraph(f'{idx}. {point.get("title", "")}', style='List Number')
                    if point.get("description"):
                        doc.add_paragraph(point["description"])

            if insights.get("framework"):
                doc.add_paragraph('核心框架:', style='Heading 3')
                for item in insights["framework"]:
                    doc.add_paragraph(item, style='List Bullet')

        # 实践行动
        doc.add_heading('✍️ 实践行动', 1)
        practice = content.get("practice", {})
        if practice:
            doc.add_heading(practice.get('title', ''), 2)
            doc.add_paragraph(practice.get('subtitle', ''))

            if practice.get("actions"):
                for idx, action in enumerate(practice["actions"], 1):
                    doc.add_heading(f'步骤 {idx}: {action.get("title", "")}', 3)
                    if action.get("description"):
                        doc.add_paragraph(action["description"])
                    if action.get("steps"):
                        for step in action["steps"]:
                            doc.add_paragraph(step, style='List Bullet')

        doc.add_page_break()

    # 我的思考
    doc.add_heading('🤔 我的思考与反思', 1)

    has_notes = False
    for idx, question in enumerate(content.get("reflection", {}).get("questions", []), 1):
        note_key = f"q{idx}"
        if notes.get(note_key):
            has_notes = True
            doc.add_heading(f'问题 {idx}', 2)
            doc.add_paragraph(question.get("text", ""))
            doc.add_paragraph('我的答案:', style='Heading 3')
            doc.add_paragraph(notes[note_key])

    if not has_notes:
        doc.add_paragraph('<i>还没有填写任何笔记。</i>')

    # 金句
    if content.get("quotes"):
        doc.add_page_break()
        doc.add_heading('💎 值得记住的话', 1)
        for quote in content["quotes"]:
            p = doc.add_paragraph(quote)
            p.italic = True

    # 保存到BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    word_bytes = buffer.getvalue()
    buffer.close()

    return word_bytes


def render_library():
    """书籍库页面 - 优化版"""
    # 滚动到顶部
    scroll_to_top()

    # 页面头部
    st.markdown("""
<div class="page-header">
    <div class="page-title">🧠 深度阅读</div>
    <div class="page-subtitle">给思考留出时间</div>
</div>
""", unsafe_allow_html=True)

    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    # ========== 控制面板：搜索、排序、视图 ==========
    st.markdown('<div class="section-block">', unsafe_allow_html=True)

    # 搜索框（保留在外面）
    search_query = st.text_input(
        "🔍 搜索书籍",
        placeholder="输入书名、作者或关键词...",
        label_visibility="visible",
        key="book_search"
    )

    st.markdown('</div>', unsafe_allow_html=True)

    # ========== 排序和视图选择（折叠） ==========
    with st.expander("🎛️ 排序与视图设置（点击展开）", expanded=False):
        # 初始化排序选项
        if 'sort_option' not in st.session_state:
            st.session_state.sort_option = "默认"

        sort_option = st.selectbox(
            "📊 排序方式",
            ["默认", "书名 A-Z", "书名 Z-A", "作者", "可阅读优先"],
            label_visibility="visible",
            key="sort_select"
        )
        st.session_state.sort_option = sort_option

    st.markdown('</div>', unsafe_allow_html=True)

    # ========== 标签筛选（折叠） ==========
    with st.expander("🏷️ 按主题筛选（点击展开）", expanded=False):
        st.markdown('<div style="display: flex; align-items: baseline; gap: 0.75rem; margin-bottom: 1rem;"><span style="font-size: 0.9rem; color: #636E72;">点击选择主题，可多选</span></div>', unsafe_allow_html=True)

        # 收集所有标签
        all_tags = set()
        for book in BOOKS_DATA:
            all_tags.update(book['tags'])

        # 标签按钮（横向排列，每行最多4个）
        tags_list = sorted(all_tags)
        num_rows = (len(tags_list) + 3) // 4  # 每行4个标签

        # 存储选中的标签
        if 'selected_tags' not in st.session_state:
            st.session_state.selected_tags = []

        # "全部"按钮和标签筛选状态显示
        col_clear, col_status = st.columns([1, 3])
        with col_clear:
            if st.button("📚 全部清除", key="tag_clear_all", use_container_width=True):
                st.session_state.selected_tags = []
                st.rerun()

        with col_status:
            if st.session_state.selected_tags:
                st.markdown(f'<div style="padding: 0.5rem; background: linear-gradient(145deg, #e8eef2 0%, #dfe6ed 100%); border-radius: 12px; text-align: center; color: #2D3436; font-size: 0.85rem; font-weight: 500;">已选: {", ".join(st.session_state.selected_tags)}</div>', unsafe_allow_html=True)
            else:
                st.markdown(f'<div style="padding: 0.5rem; background: #F0F3F5; border-radius: 12px; text-align: center; color: #636E72; font-size: 0.85rem;">未选择任何主题</div>', unsafe_allow_html=True)

        st.markdown('<br>', unsafe_allow_html=True)

        # 标签按钮（4列网格）
        for row in range(num_rows):
            tag_cols = st.columns(4)
            for col in range(4):
                tag_idx = row * 4 + col
                if tag_idx < len(tags_list):
                    tag = tags_list[tag_idx]
                    with tag_cols[col]:
                        is_selected = tag in st.session_state.selected_tags
                        if is_selected:
                            st.markdown(f"""
<div style="padding: 0.5rem; background: linear-gradient(145deg, #667eea 0%, #764ba2 100%); border-radius: 12px; text-align: center; color: #ffffff; font-size: 0.8rem; font-weight: 600; cursor: pointer; border: 2px solid rgba(102, 126, 234, 0.3);">
    ✓ {tag}
</div>
""", unsafe_allow_html=True)
                            if st.button(f"取消 {tag}", key=f"tag_{tag}_off", use_container_width=True):
                                st.session_state.selected_tags.remove(tag)
                                st.rerun()
                        else:
                            st.markdown(f"""
<div style="padding: 0.5rem; background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%); border-radius: 12px; text-align: center; color: #2D3436; font-size: 0.8rem; font-weight: 500; cursor: pointer; border: 1px solid rgba(102, 126, 234, 0.1);">
    {tag}
</div>
""", unsafe_allow_html=True)
                        if st.button(f"选择 {tag}", key=f"tag_{tag}_on", use_container_width=True):
                            st.session_state.selected_tags.append(tag)
                            st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)

    # ========== 视图切换（折叠） ==========
    with st.expander("🎨 视图切换（点击展开）", expanded=False):
        if 'view_mode' not in st.session_state:
            st.session_state.view_mode = "grid_2"  # 默认2列网格

        view_cols = st.columns(3)
        with view_cols[0]:
            if st.button("2列", key="view_2col", use_container_width=True, type="primary" if st.session_state.view_mode == "grid_2" else "secondary"):
                st.session_state.view_mode = "grid_2"
                st.rerun()
        with view_cols[1]:
            if st.button("3列", key="view_3col", use_container_width=True, type="primary" if st.session_state.view_mode == "grid_3" else "secondary"):
                st.session_state.view_mode = "grid_3"
                st.rerun()
        with view_cols[2]:
            if st.button("列表", key="view_list", use_container_width=True, type="primary" if st.session_state.view_mode == "list" else "secondary"):
                st.session_state.view_mode = "list"
                st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)
    # ============================================

    # 筛选书籍 - 结合搜索和标签
    filtered_books = BOOKS_DATA.copy()

    # 先按标签筛选
    if st.session_state.selected_tags:
        filtered_books = [
            book for book in filtered_books
            if any(tag in book['tags'] for tag in st.session_state.selected_tags)
        ]

    # 再按搜索关键词筛选
    if search_query and search_query.strip():
        search_query = search_query.lower().strip()
        filtered_books = [
            book for book in filtered_books
            if (search_query in book['title'].lower() or
                search_query in book['author'].lower() or
                search_query in book['description'].lower() or
                any(search_query in tag.lower() for tag in book['tags']))
        ]

    # 排序
    if sort_option == "书名 A-Z":
        filtered_books.sort(key=lambda x: x['title'].lower())
    elif sort_option == "书名 Z-A":
        filtered_books.sort(key=lambda x: x['title'].lower(), reverse=True)
    elif sort_option == "作者":
        filtered_books.sort(key=lambda x: x['author'].lower())
    elif sort_option == "可阅读优先":
        filtered_books.sort(key=lambda x: not x['available'])

    # 显示结果数量
    st.markdown(f'<div style="text-align: center; color: #636E72; font-size: 0.8rem; margin: 1rem 0; padding: 0.5rem; background: linear-gradient(145deg, #f8f9fa 0%, #e8eef2 100%); border-radius: 12px; display: inline-block; width: 100%; box-sizing: border-box;">📚 显示 {len(filtered_books)} 本书</div>', unsafe_allow_html=True)
    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)
    # ============================================

    # 根据视图模式渲染书籍
    if st.session_state.view_mode == "grid_2":
        # 2列网格
        for i in range(0, len(filtered_books), 2):
            book1 = filtered_books[i]
            book2 = filtered_books[i + 1] if i + 1 < len(filtered_books) else None

            if book2:
                col1, col2 = st.columns(2)
                with col1:
                    render_book_card(book1)
                with col2:
                    render_book_card(book2)
            else:
                render_book_card(book1, center=True)

    elif st.session_state.view_mode == "grid_3":
        # 3列网格
        for i in range(0, len(filtered_books), 3):
            books_in_row = filtered_books[i:i+3]
            cols = st.columns(len(books_in_row))
            for col, book in zip(cols, books_in_row):
                with col:
                    render_book_card(book)

    else:  # list view
        # 列表视图
        for book in filtered_books:
            render_book_card_list(book)


def render_book_card(book, center=False):
    """渲染单本书籍卡片"""
    card_class = "available" if book["available"] else "unavailable"

    # 初始化收藏列表
    if 'favorite_books' not in st.session_state:
        st.session_state.favorite_books = []

    # 检查是否已收藏
    is_favorite = book['title'] in st.session_state.favorite_books
    fav_emoji = "❤️" if is_favorite else "🤍"
    fav_title = "取消收藏" if is_favorite else "收藏"

    # 创建一个容器来包含卡片和按钮
    st.markdown(f'<div style="position: relative;">', unsafe_allow_html=True)

    # 书籍卡片
    st.markdown(f"""
<div class="book-card {card_class}">
    <div class="book-cover">{book['emoji']}</div>
    <div class="book-info">
        <div>
            <div class="book-title">{book['title']}</div>
            <div class="book-author">{book['author']}</div>
            <div class="book-description">{book['description']}</div>
            <div class="tag-container">
                {' '.join([f'<span class="tag {("highlight" if book["available"] else "")}">{tag}</span>' for tag in book['tags']])}
            </div>
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

    # 按钮行 - 收藏和阅读
    if book["available"]:
        col_fav, col_read, col_del = st.columns([1, 4, 1])

        with col_fav:
            if st.button(fav_emoji, key=f"fav_{book['title']}", help=fav_title):
                if is_favorite:
                    st.session_state.favorite_books.remove(book['title'])
                else:
                    st.session_state.favorite_books.append(book['title'])
                st.rerun()

        with col_read:
            if st.button(f"📖 开始阅读", key=f"read_{book['title']}", use_container_width=True):
                st.session_state.page_rerun += 1
                st.session_state.current_book = book['title']
                st.session_state.current_content = get_book_content(book['title'])
                st.session_state.current_section = "intro"
                st.rerun()

        with col_del:
            if st.button("🗑️", key=f"del_{book['title']}", help="删除书籍"):
                # 确认删除
                if f"confirm_del_{book['title']}" not in st.session_state:
                    st.session_state[f"confirm_del_{book['title']}"] = False

                if st.session_state[f"confirm_del_{book['title']}"]:
                    # 执行删除 - 从BOOKS_DATA中移除（通过标记为unavailable）
                    for b in BOOKS_DATA:
                        if b['title'] == book['title']:
                            b['available'] = False
                            break
                    st.success(f"已删除《{book['title']}》")
                    st.rerun()
                else:
                    st.session_state[f"confirm_del_{book['title']}"] = True
                    st.rerun()

                # 如果在确认状态，显示取消按钮
                if st.session_state[f"confirm_del_{book['title']}"]:
                    if st.button("取消", key=f"cancel_del_{book['title']}", use_container_width=True):
                        st.session_state[f"confirm_del_{book['title']}"] = False
                        st.rerun()
    else:
        st.markdown(f'<div style="text-align: center; color: #636E72; font-size: 0.75rem; font-style: italic; margin-top: 0.5rem;">即将推出</div>', unsafe_allow_html=True)


def render_book_card_list(book):
    """列表视图的书籍卡片"""
    card_class = "available" if book["available"] else "unavailable"

    # 初始化收藏列表
    if 'favorite_books' not in st.session_state:
        st.session_state.favorite_books = []

    # 检查是否已收藏
    is_favorite = book['title'] in st.session_state.favorite_books
    fav_emoji = "❤️" if is_favorite else "🤍"

    # 列表视图 - 横向布局
    st.markdown(f"""
<div style="background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%);
     border-radius: 16px; padding: 1.25rem; margin-bottom: 1rem;
     border: 1px solid rgba(102, 126, 234, 0.1);
     box-shadow: 0 2px 8px rgba(0, 0, 0, 0.05);
     transition: all 0.3s ease; display: flex; gap: 1.5rem; align-items: center;">
    <div style="font-size: 3.5rem; text-align: center; min-width: 100px;">
        {book['emoji']}
    </div>
    <div style="flex: 1;">
        <div style="font-family: 'Noto Serif SC', serif; font-size: 1.2rem; font-weight: 700; color: #2D3436; margin-bottom: 0.3rem;">
            {book['title']}
        </div>
        <div style="font-family: 'Inter', sans-serif; font-size: 0.8rem; color: #667eea; font-weight: 600; margin-bottom: 0.5rem; text-transform: uppercase; letter-spacing: 0.05em;">
            {book['author']}
        </div>
        <div style="font-family: 'Noto Serif SC', serif; font-size: 0.85rem; color: #636E72; line-height: 1.5; margin-bottom: 0.75rem;">
            {book['description']}
        </div>
        <div style="display: flex; gap: 0.4rem; flex-wrap: wrap;">
            {(' '.join([f'<span class="tag {("highlight" if book["available"] else "")}">{tag}</span>' for tag in book['tags']]))}
        </div>
    </div>
</div>
""", unsafe_allow_html=True)

    # 按钮行
    col_fav, col_read = st.columns([1, 4])

    with col_fav:
        if st.button(fav_emoji, key=f"fav_list_{book['title']}", help="取消收藏" if is_favorite else "收藏"):
            if is_favorite:
                st.session_state.favorite_books.remove(book['title'])
            else:
                st.session_state.favorite_books.append(book['title'])
            st.rerun()

    with col_read:
        if book["available"]:
            if st.button(f"📖 开始阅读", key=f"read_list_{book['title']}", use_container_width=True, type="primary"):
                st.session_state.page_rerun += 1
                st.session_state.current_book = book['title']
                st.session_state.current_content = get_book_content(book['title'])
                st.session_state.current_section = "intro"
                st.rerun()
        else:
            st.markdown(f'<div style="text-align: center; color: #636E72; font-size: 0.75rem; font-style: italic; padding: 0.5rem;">即将推出</div>', unsafe_allow_html=True)


def render_introduction(content):
    """引言页 - 优雅简洁版"""
    intro = content["introduction"]

    # 滚动到顶部
    scroll_to_top()

    # 顶部导航
    col1, col2, col3 = st.columns([1, 4, 1])
    with col1:
        if st.button("← 返回", key="intro_back_library"):
            st.session_state.page_rerun += 1
            st.session_state.current_book = None
            st.session_state.current_content = None
            st.session_state.current_section = "library"
            st.rerun()

    # 阅读进度条
    render_progress_bar("intro")

    # 页面标题
    st.markdown(f'<div class="section-title">{intro["title"]}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="section-subtitle">{intro["subtitle"]}</div>', unsafe_allow_html=True)

    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    # 内容分区
    st.markdown('<div class="section-block">', unsafe_allow_html=True)

    st.markdown('<div class="subsection-header">为什么要读这本书</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="content-block"><div class="body-text">{intro["why_read"]}</div></div>', unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

    # 阅读前提问
    st.markdown('<div class="section-block">', unsafe_allow_html=True)
    st.markdown('<div class="subsection-header">阅读前，先问问自己</div>', unsafe_allow_html=True)

    for i, question in enumerate(intro["pre_questions"], 1):
        st.markdown(f"""
<div class="question-block">
    <div class="question-label">问题 {i}</div>
    <div class="question-text">{question}</div>
</div>
""", unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

    # 开始阅读
    st.markdown('<div class="nav-container">', unsafe_allow_html=True)
    if st.button("📖 开始深入阅读", key="intro_start", use_container_width=True, type="primary"):
        st.session_state.page_rerun += 1
        st.session_state.current_section = "insights"
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)


def render_insights(content):
    """核心洞察页 - 优雅简洁版"""
    core = content["core_thinking"]

    # 滚动到顶部
    scroll_to_top()

    # 顶部导航
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("← 引言", key="insights_back_intro"):
            st.session_state.page_rerun += 1
            st.session_state.current_section = "intro"
            st.rerun()

    st.markdown(f'<div class="section-title">{core["title"]}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="section-subtitle">{core["subtitle"]}</div>', unsafe_allow_html=True)

    # 阅读进度条
    render_progress_bar("insights")

    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    # 每个洞察
    for idx, insight in enumerate(core["insights"], 1):
        # 简约线条编号
        st.markdown(f'<div class="insight-number">— {idx:02d} —</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="section-title">{insight["title"]}</div>', unsafe_allow_html=True)

        # 核心观点 - 简洁优雅
        st.markdown('<div class="core-idea-box">', unsafe_allow_html=True)
        core_idea = insight["core_idea"].strip()
        core_idea = '\n\n'.join(line.strip() for line in core_idea.split('\n') if line.strip())
        st.markdown(f'<div class="core-idea-text">{core_idea}</div>', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

        # 分区内容
        st.markdown('<div class="section-block">', unsafe_allow_html=True)

        st.markdown('<div class="subsection-header">为什么这很重要</div>', unsafe_allow_html=True)
        st.markdown(f'<div class="content-block"><div class="body-text">{insight["why_matters"]}</div></div>', unsafe_allow_html=True)

        if insight.get("example"):
            st.markdown('<div class="subsection-header">现实中的样子</div>', unsafe_allow_html=True)
            st.markdown(f'<div class="content-block highlight"><div class="body-text">{insight["example"]}</div></div>', unsafe_allow_html=True)

        if insight.get("question"):
            st.markdown(f"""
<div class="question-block">
    <div class="question-label">想一想</div>
    <div class="question-text">{insight["question"]}</div>
</div>
""", unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)

        if idx < len(core["insights"]):
            st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    # 底部导航
    st.markdown('<div class="nav-container">', unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("← 引言", key="insights_bottom_back"):
            st.session_state.page_rerun += 1
            st.session_state.current_section = "intro"
            st.rerun()

    with col3:
        if st.button("实践 →", key="insights_to_practice"):
            st.session_state.page_rerun += 1
            st.session_state.current_section = "practice"
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)


def render_practice(content):
    """实践页 - 优化版"""
    practice = content["practice"]

    # 滚动到顶部
    scroll_to_top()

    # 顶部导航
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("← 洞察", key="practice_back_insights"):
            st.session_state.page_rerun += 1
            st.session_state.current_section = "insights"
            st.rerun()

    st.markdown(f'<div class="section-title">{practice["title"]}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="section-subtitle">{practice["subtitle"]}</div>', unsafe_allow_html=True)

    # 阅读进度条
    render_progress_bar("practice")

    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    # 实践步骤 - 优化布局
    for idx, item in enumerate(practice["actions"], 1):
        st.markdown('<div class="section-block">', unsafe_allow_html=True)

        # 添加编号徽章
        st.markdown(f"""
<div style="display: flex; align-items: center; gap: 0.75rem; margin-bottom: 1rem;">
    <div style="background: linear-gradient(145deg, #667eea 0%, #764ba2 100%);
         color: #ffffff; width: 32px; height: 32px; border-radius: 50%;
         display: flex; align-items: center; justify-content: center;
         font-weight: 700; font-size: 0.9rem; flex-shrink: 0;">
        {idx}
    </div>
    <div class="subsection-header" style="margin: 0;">{item["title"]}</div>
</div>
""", unsafe_allow_html=True)

        st.markdown(f'<div class="body-text" style="margin-bottom: 1.25rem; margin-left: 2.5rem;">{item["description"]}</div>', unsafe_allow_html=True)

        if item.get("steps"):
            st.markdown('<div class="step-list" style="margin-left: 2.5rem;">', unsafe_allow_html=True)
            for step in item["steps"]:
                st.markdown(f'<div class="step-item"><div class="step-number">✓</div><div class="step-text">{step}</div></div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('</div>', unsafe_allow_html=True)

    # ========== 30天实践计划入口 ==========
    if content["title"] in PRACTICE_TASKS:
        st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

        # 检查是否已经开始
        book_title = content["title"]
        has_started = book_title in st.session_state.practice_tracker

        # 优雅的卡片设计
        st.markdown(f"""
<div style="background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%);
     border-radius: 16px; padding: 2rem; margin: 2rem 0;
     border: 2px solid rgba(102, 126, 234, 0.15);
     box-shadow: 0 4px 12px rgba(102, 126, 234, 0.1);
     text-align: center; position: relative; overflow: hidden;">
    <div style="position: absolute; top: 0; left: 0; right: 0; height: 3px;
         background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);"></div>
    <div style="font-size: 1.5rem; font-weight: 700; color: #2D3436; margin-bottom: 0.5rem;">
        🎯 30天实践计划
    </div>
    <div style="font-size: 0.95rem; color: #636E72; margin-bottom: 1.5rem; line-height: 1.6;">
        不只是"知道"，更是"做到"<br/>
        每日任务 · 可追踪 · 成就系统
    </div>
</div>
""", unsafe_allow_html=True)

        if has_started:
            # 已开始，显示继续按钮
            if st.button("📊 继续我的实践计划", key="continue_practice", use_container_width=True, type="primary"):
                st.session_state.current_section = "practice_tasks"
                st.rerun()
        else:
            # 未开始，显示开始按钮
            if st.button("🚀 开始30天挑战", key="start_practice", use_container_width=True, type="primary"):
                # 初始化追踪数据
                st.session_state.practice_tracker[book_title] = {
                    "start_date": datetime.now().strftime("%Y-%m-%d"),
                    "current_day": 1,
                    "completed_days": {},
                    "badges": []
                }
                st.session_state.current_section = "practice_tasks"
                st.rerun()
    # =========================================

    # 底部导航
    st.markdown('<div class="nav-container">', unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("← 洞察", key="practice_bottom_back"):
            st.session_state.page_rerun += 1
            st.session_state.current_section = "insights"
            st.rerun()

    with col3:
        if st.button("反思 →", key="practice_to_reflection"):
            st.session_state.page_rerun += 1
            st.session_state.current_section = "reflection"
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)


def generate_30day_poster(user_habits, completed_count, consecutive, total_progress, habits_completion, book_title, content):
    """生成30天成就海报"""
    # 计算每个习惯的完成情况
    habits_stats = []
    for habit in user_habits[:3]:
        habit_completed = sum(1 for day_data in habits_completion.values() if day_data.get(habit, False))
        habit_percentage = int((habit_completed / 30) * 100)
        habits_stats.append({"name": habit, "completed": habit_completed, "percentage": habit_percentage})

    # 生成热力图
    heatmap_grid = ""
    for week in range(4):
        heatmap_grid += "<div style='display: flex; gap: 4px; margin-bottom: 4px;'>"
        for day in range(7):
            day_num = week * 7 + day + 1
            if day_num > 30:
                break

            day_str = str(day_num)
            day_habits = habits_completion.get(day_str, {})
            completed_today = sum(day_habits.values())
            percentage = int((completed_today / 3) * 100)

            if percentage == 100:
                color = "#4CAF50"
            elif percentage >= 66:
                color = "#8BC34A"
            elif percentage >= 33:
                color = "#FFC107"
            else:
                color = "#E0E0E0"

            heatmap_grid += f"<div style='width: 28px; height: 28px; background: {color}; border-radius: 4px;'></div>"
        heatmap_grid += "</div>"

    poster_html = f"""
<div style="width: 100%; max-width: 600px; margin: 2rem auto; padding: 0; background: #ffffff; border-radius: 24px; box-shadow: 0 20px 60px rgba(0, 0, 0, 0.15); overflow: hidden;">
    <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 3rem 2rem; text-align: center; position: relative;">
        <div style="font-size: 4rem; margin-bottom: 1rem;">🎉</div>
        <div style="font-size: 2rem; color: #ffffff; font-weight: 700; margin-bottom: 0.5rem;">30天挑战成功！</div>
        <div style="font-size: 1rem; color: rgba(255, 255, 255, 0.9); margin-bottom: 1rem;">{book_title} · 微习惯养成计划</div>
        <div style="display: flex; justify-content: center; gap: 2rem;">
            <div>
                <div style="font-size: 2rem; color: #ffffff; font-weight: 700;">{completed_count}</div>
                <div style="font-size: 0.75rem; color: rgba(255, 255, 255, 0.8);">总打卡</div>
            </div>
            <div>
                <div style="font-size: 2rem; color: #ffffff; font-weight: 700;">{consecutive}</div>
                <div style="font-size: 0.75rem; color: rgba(255, 255, 255, 0.8);">连续天数</div>
            </div>
            <div>
                <div style="font-size: 2rem; color: #ffffff; font-weight: 700;">{total_progress}%</div>
                <div style="font-size: 0.75rem; color: rgba(255, 255, 255, 0.8);">完成度</div>
            </div>
        </div>
    </div>

    <div style="padding: 2rem;">
        <div style="margin-bottom: 2rem;">
            <div style="font-size: 1rem; color: #2D3436; font-weight: 600; margin-bottom: 1rem; text-align: center;">📊 三个习惯完成情况</div>
            <div style="display: flex; gap: 1rem; justify-content: center;">
                {''.join([f"""
                <div style="flex: 1; text-align: center; padding: 1rem; background: #F8F9FA; border-radius: 12px;">
                    <div style="font-size: 0.8rem; color: #636E72; margin-bottom: 0.5rem;">{stat['name'][:15]}...</div>
                    <div style="font-size: 1.5rem; color: #667eea; font-weight: 700;">{stat['completed']}/30</div>
                    <div style="font-size: 0.9rem; color: #764ba2;">{stat['percentage']}%</div>
                </div>
                """ for stat in habits_stats])}
            </div>
        </div>

        <div style="margin-bottom: 2rem;">
            <div style="font-size: 1rem; color: #2D3436; font-weight: 600; margin-bottom: 1rem; text-align: center;">📅 30天打卡热力图</div>
            <div style="display: flex; justify-content: center;">
                {heatmap_grid}
            </div>
            <div style="display: flex; justify-content: center; gap: 1rem; margin-top: 1rem; font-size: 0.75rem; color: #636E72;">
                <div style="display: flex; align-items: center; gap: 0.25rem;"><div style="width: 16px; height: 16px; background: #4CAF50; border-radius: 3px;"></div> 全完成</div>
                <div style="display: flex; align-items: center; gap: 0.25rem;"><div style="width: 16px; height: 16px; background: #8BC34A; border-radius: 3px;"></div> 良好</div>
                <div style="display: flex; align-items: center; gap: 0.25rem;"><div style="width: 16px; height: 16px; background: #FFC107; border-radius: 3px;"></div> 一般</div>
                <div style="display: flex; align-items: center; gap: 0.25rem;"><div style="width: 16px; height: 16px; background: #E0E0E0; border-radius: 3px;"></div> 未完成</div>
            </div>
        </div>

        <div style="text-align: center; padding-top: 1.5rem; border-top: 1px solid #E8EEF2;">
            <div style="display: inline-flex; align-items: center; gap: 0.5rem; padding: 0.75rem 1.5rem; background: rgba(102, 126, 234, 0.1); border-radius: 25px; color: #667eea; font-weight: 600; font-size: 0.9rem;">
                <span>🧠</span><span>DeepRead 深读</span>
            </div>
            <div style="font-size: 0.75rem; color: #636E72; margin-top: 0.75rem; font-style: italic;">微习惯 · 大改变</div>
        </div>
    </div>
</div>

<div style="text-align: center; margin: 2rem 0; color: #636E72; font-size: 0.85rem;">
<strong>💡 如何分享：</strong><br/>1. 在电脑上：截图后保存图片<br/>2. 在手机上：长按海报区域保存图片<br/>3. 分享到朋友圈、小红书、微博等社交平台
</div>
"""
    st.markdown(poster_html, unsafe_allow_html=True)


def render_practice_tasks(content):
    """30天实践任务追踪页 - 轻松版本"""
    book_title = content["title"]

    # 检查是否有实践任务
    if book_title not in PRACTICE_TASKS:
        st.error("此书籍暂无实践任务")
        return

    practice_data = PRACTICE_TASKS[book_title]

    # 滚动到顶部
    scroll_to_top()

    # 顶部导航
    col1, col2 = st.columns([1, 4])
    with col1:
        if st.button("← 返回", key="practice_tasks_back"):
            st.session_state.current_section = "practice"
            st.rerun()

    # 获取用户进度数据
    tracker = st.session_state.practice_tracker.get(book_title, {})
    # 新的数据结构：记录每天每个习惯的完成情况
    # habits_completion: { "1": {"habit1": true, "habit2": false, "habit3": true}, ... }
    habits_completion = tracker.get("habits_completion", {})

    # 计算总完成度：3个习惯 × 30天 = 90个可能的打卡
    total_slots = 3 * 30
    completed_count = 0
    for day_data in habits_completion.values():
        completed_count += sum(day_data.values())
    total_progress = int((completed_count / total_slots) * 100)

    # 计算连续天数（以天为单位，当天3个习惯都完成才算）
    consecutive = 0
    for day in range(1, 31):
        day_str = str(day)
        if day_str in habits_completion:
            day_habits = habits_completion[day_str]
            if len(day_habits) == 3 and all(day_habits.values()):
                consecutive += 1
            else:
                break
        else:
            break

    # 让用户选择要养成的3个微习惯
    user_habits = tracker.get("user_habits", [])

    # 检查是否完成了30天挑战（必须100%完成）
    if total_progress >= 100 and user_habits:
        # 触发成就检查
        check_and_unlock_achievements()

        # 显示30天完成庆祝页面
        st.balloons()

        st.markdown(f"""<div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 3rem 2rem; border-radius: 24px; margin-bottom: 2rem; text-align: center; position: relative; overflow: hidden;">
    <div style="position: absolute; top: -80px; right: -80px; width: 200px; height: 200px; background: rgba(255, 255, 255, 0.1); border-radius: 50%;"></div>
    <div style="position: absolute; bottom: -60px; left: -60px; width: 150px; height: 150px; background: rgba(255, 255, 255, 0.08); border-radius: 50%;"></div>

    <div style="position: relative; z-index: 2;">
        <div style="font-size: 5rem; margin-bottom: 1rem;">🎉</div>
        <div style="font-size: 2.5rem; font-weight: 700; color: #ffffff; margin-bottom: 1rem;">恭喜你完成了30天挑战！</div>
        <div style="font-size: 1.2rem; color: rgba(255, 255, 255, 0.9); margin-bottom: 2rem;">
            你用坚持证明了：微习惯的力量是巨大的
        </div>

        <div style="display: flex; justify-content: center; gap: 3rem; margin-top: 2rem; flex-wrap: wrap;">
            <div style="text-align: center;">
                <div style="font-size: 3rem; font-weight: 700; color: #ffffff;">{completed_count}</div>
                <div style="font-size: 0.9rem; color: rgba(255, 255, 255, 0.8); margin-top: 0.5rem;">总打卡数</div>
            </div>
            <div style="text-align: center;">
                <div style="font-size: 3rem; font-weight: 700; color: #ffffff;">{consecutive}</div>
                <div style="font-size: 0.9rem; color: rgba(255, 255, 255, 0.8); margin-top: 0.5rem;">最长连续天数</div>
            </div>
            <div style="text-align: center;">
                <div style="font-size: 3rem; font-weight: 700; color: #ffffff;">{total_progress}%</div>
                <div style="font-size: 0.9rem; color: rgba(255, 255, 255, 0.8); margin-top: 0.5rem;">完成度</div>
            </div>
        </div>
    </div>
</div>""", unsafe_allow_html=True)

        # 30天打卡热力图
        st.markdown("### 📅 30天打卡记录")
        st.markdown("**每一天的坚持，都是向更好的自己迈进**", unsafe_allow_html=True)

        heatmap_data = []
        for day in range(1, 31):
            day_str = str(day)
            day_habits = habits_completion.get(day_str, {})
            completed_today = sum(day_habits.values())
            percentage = int((completed_today / 3) * 100)

            if percentage == 100:
                color = "#4CAF50"
                emoji = "🌟"
            elif percentage >= 66:
                color = "#8BC34A"
                emoji = "✓"
            elif percentage >= 33:
                color = "#FFC107"
                emoji = "○"
            else:
                color = "#E0E0E0"
                emoji = "·"

            heatmap_data.append({"day": day, "color": color, "emoji": emoji, "percentage": percentage})

        # 显示热力图
        weeks_grid = [heatmap_data[i:i+7] for i in range(0, 30, 7)]

        for week_idx, week_data in enumerate(weeks_grid, 1):
            cols = st.columns(7)
            for col_idx, day_data in enumerate(week_data):
                with cols[col_idx]:
                    st.markdown(f"""
<div style="background: {day_data['color']}; padding: 1rem 0.5rem; border-radius: 12px; text-align: center; min-height: 80px; display: flex; flex-direction: column; justify-content: center; align-items: center;">
    <div style="font-size: 1.5rem;">{day_data['emoji']}</div>
    <div style="font-size: 0.75rem; color: rgba(0,0,0,0.6); margin-top: 0.25rem;">第{day_data['day']}天</div>
</div>
""", unsafe_allow_html=True)

        # 反思区域
        st.markdown("---")
        st.markdown("### 💭 写下你的30天感悟")
        st.markdown("**回顾这段旅程，记录你的成长和变化**", unsafe_allow_html=True)

        reflection_key = f"30day_reflection_{book_title}"
        saved_reflection = st.session_state.get(reflection_key, "")

        reflection = st.text_area(
            "我的30天心得...",
            value=saved_reflection,
            placeholder="这30天给我带来的改变是...",
            height=150,
            key="reflection_input"
        )

        col_save_ref1, col_save_ref2, col_save_ref3 = st.columns([1, 2, 1])
        with col_save_ref2:
            if st.button("💾 保存我的感悟", use_container_width=True, type="primary"):
                st.session_state[reflection_key] = reflection
                st.success("✅ 感悟已保存！感谢你的坚持！")
                st.balloons()

        # 下一步建议
        st.markdown("---")
        st.markdown("### 🚀 接下来做什么？")

        suggestion_cols = st.columns(3)
        with suggestion_cols[0]:
            st.markdown("""
<div style="background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%); padding: 1.5rem; border-radius: 16px; border: 1px solid rgba(102, 126, 234, 0.1); text-align: center; min-height: 220px; display: flex; flex-direction: column; justify-content: center;">
    <div style="font-size: 3rem; margin-bottom: 1rem;">🔄</div>
    <div style="font-size: 1.1rem; font-weight: 600; color: #2D3436; margin-bottom: 0.5rem;">继续这3个习惯</div>
    <div style="font-size: 0.85rem; color: #636E72;">习惯已经养成，继续巩固让它成为自动行为</div>
</div>
""", unsafe_allow_html=True)

        with suggestion_cols[1]:
            st.markdown("""
<div style="background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%); padding: 1.5rem; border-radius: 16px; border: 1px solid rgba(102, 126, 234, 0.1); text-align: center; min-height: 220px; display: flex; flex-direction: column; justify-content: center;">
    <div style="font-size: 3rem; margin-bottom: 1rem;">⬆️</div>
    <div style="font-size: 1.1rem; font-weight: 600; color: #2D3436; margin-bottom: 0.5rem;">提升难度</div>
    <div style="font-size: 0.85rem; color: #636E72;">增加时长或强度，让习惯更有挑战性</div>
</div>
""", unsafe_allow_html=True)

        with suggestion_cols[2]:
            st.markdown("""
<div style="background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%); padding: 1.5rem; border-radius: 16px; border: 1px solid rgba(102, 126, 234, 0.1); text-align: center; min-height: 220px; display: flex; flex-direction: column; justify-content: center;">
    <div style="font-size: 3rem; margin-bottom: 1rem;">➕</div>
    <div style="font-size: 1.1rem; font-weight: 600; color: #2D3436; margin-bottom: 0.5rem;">培养新习惯</div>
    <div style="font-size: 0.85rem; color: #636E72;">开启新的30天挑战，养成更多好习惯</div>
</div>
""", unsafe_allow_html=True)

        # 生成30天成就海报按钮
        st.markdown("---")
        st.markdown("### 🖼️ 分享你的成就")
        if st.button("🎨 生成30天成就海报", use_container_width=True, type="primary"):
            generate_30day_poster(user_habits, completed_count, consecutive, total_progress, habits_completion, book_title, content)

        # 重新开始按钮
        if st.button("🔄 重新开始30天挑战", key="restart_30days"):
            if book_title in st.session_state.practice_tracker:
                del st.session_state.practice_tracker[book_title]["habits_completion"]
            st.session_state.selected_week = 1
            st.rerun()

        return

    # 让用户选择要养成的3个微习惯
    user_habits = tracker.get("user_habits", [])
    if not user_habits or len(user_habits) == 0:
        st.markdown("### 🌱 设定你的三个微习惯")
        st.markdown("**选择3个简单的微习惯，用30天养成它们！**", unsafe_allow_html=True)

        st.markdown('<br>', unsafe_allow_html=True)

        # 初始化session state
        if "selected_habits" not in st.session_state:
            st.session_state.selected_habits = []

        # 预设选项
        habit_options = [
            "📚 每天阅读10分钟",
            "🧘 每天冥想5分钟",
            "🏃 每天运动15分钟",
            "✍️ 每天写日记",
            "⏰ 每天早起10分钟",
            "💧 每天喝8杯水",
            "📖 每天背诵5个单词",
            "🏠 每天整理房间10分钟",
            "📵 每天不看手机1小时",
            "🙏 每天感恩3件事",
            "🥗 每天吃一份水果",
            "🌙 每天早睡10分钟",
            "💪 每天做10个深蹲",
            "🎨 每天画画5分钟",
            "🎵 每天听一首新歌"
        ]

        # 显示当前已选择的习惯
        if st.session_state.selected_habits:
            st.markdown("### ✅ 已选择的习惯")
            for idx, habit in enumerate(st.session_state.selected_habits, 1):
                st.markdown(f"**{idx}. {habit}**")

            st.markdown(f"**已选择 {len(st.session_state.selected_habits)}/3 个习惯**")

            if len(st.session_state.selected_habits) >= 3:
                st.success("🎉 已完成选择！")
                if st.button("✓ 确认并开始", type="primary", use_container_width=True):
                    if book_title not in st.session_state.practice_tracker:
                        st.session_state.practice_tracker[book_title] = {"completed_days": {}}
                    st.session_state.practice_tracker[book_title]["user_habits"] = st.session_state.selected_habits[:3]
                    st.success(f"太棒了！你要培养的三个习惯是：**{', '.join(st.session_state.selected_habits[:3])}**")
                    st.rerun()
            else:
                st.info(f"再选择 **{3 - len(st.session_state.selected_habits)}** 个习惯即可开始")

            st.markdown("---")

        st.markdown("### 📋 可选的微习惯")
        st.markdown("点击选择你想要养成的习惯：", unsafe_allow_html=True)

        # 使用3列布局显示选项
        cols = st.columns(3)
        for idx, habit in enumerate(habit_options):
            col_idx = idx % 3
            with cols[col_idx]:
                is_selected = habit in st.session_state.selected_habits
                button_label = f"✓ {habit}" if is_selected else habit

                if st.button(button_label, key=f"habit_{idx}", use_container_width=True):
                    if is_selected:
                        st.session_state.selected_habits.remove(habit)
                    else:
                        if len(st.session_state.selected_habits) < 3:
                            st.session_state.selected_habits.append(habit)
                        else:
                            st.warning("最多只能选择3个习惯，请先取消一个")
                    st.rerun()

        # 自定义输入
        st.markdown("---")
        st.markdown("### ✏️ 或者自定义你的习惯")
        custom_habit = st.text_input(
            "输入自定义习惯（添加emoji会更生动哦）",
            placeholder="例如：🎯 每天练习投篮10次",
            key="custom_habit_input"
        )

        col_cust1, col_cust2, col_cust3 = st.columns([1, 1, 1])
        with col_cust2:
            if st.button("添加自定义习惯", use_container_width=True):
                if custom_habit.strip():
                    if len(st.session_state.selected_habits) < 3:
                        st.session_state.selected_habits.append(custom_habit.strip())
                        st.success(f"已添加：**{custom_habit.strip()}**")
                        st.rerun()
                    else:
                        st.warning("最多只能选择3个习惯，请先取消一个")
                else:
                    st.warning("请输入一个习惯")

        st.markdown('<br>', unsafe_allow_html=True)
        return

    # 显示已选择的三个习惯
    st.markdown("### 🎯 正在培养的三个习惯")
    habits_display = st.columns(3)
    for idx, habit in enumerate(user_habits[:3]):
        with habits_display[idx]:
            st.markdown(f"""
            <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 1.5rem; border-radius: 16px; text-align: center; color: white; box-shadow: 0 8px 16px rgba(102, 126, 234, 0.2);">
                <div style="font-size: 1.1rem; font-weight: 600;">{habit}</div>
            </div>
            """, unsafe_allow_html=True)

    # 提供修改习惯的选项
    with st.expander("✏️ 想要更换习惯？"):
        st.write("**当前习惯：**", ", ".join(user_habits[:3]))
        st.warning("更换习惯会重新开始30天挑战哦！")
        if st.button("重新选择习惯", key="change_habits"):
            if book_title in st.session_state.practice_tracker:
                del st.session_state.practice_tracker[book_title]["user_habits"]
            st.session_state.selected_habits = []
            st.rerun()

    st.markdown("---")

    # 温暖的进度卡片 - 简洁清晰的设计

    col_a, col_b, col_c = st.columns([1, 2, 1])
    with col_a:
        st.metric("连续坚持", f"{consecutive}天")
    with col_c:
        st.metric("完成度", f"{total_progress}%")

    st.markdown(f"**{practice_data['subtitle']}**", unsafe_allow_html=True)
    st.progress(total_progress / 100)
    st.markdown(f"### {completed_count}/{total_slots}个习惯完成", unsafe_allow_html=True)

    # 周选择器 - 轻松的Tab风格
    weeks = [
        {"key": "week_1", "label": "第1周", "emoji": "🌱", "desc": "启动 (1-7天)"},
        {"key": "week_2", "label": "第2周", "emoji": "🌿", "desc": "稳定 (8-14天)"},
        {"key": "week_3", "label": "第3周", "emoji": "🌳", "desc": "成长 (15-21天)"},
        {"key": "week_4", "label": "第4周", "emoji": "🏆", "desc": "达成 (22-30天)"}
    ]

    # 使用session state存储当前选择的周
    if "selected_week" not in st.session_state:
        st.session_state.selected_week = 1

    # 周选择按钮
    st.markdown('<div style="margin: 1.5rem 0;">', unsafe_allow_html=True)
    st.markdown('<div style="font-size: 1rem; color: #636E72; margin-bottom: 1rem; text-align: center;">💫 点击切换到其他周</div>', unsafe_allow_html=True)
    week_cols = st.columns(4)

    for idx, week_info in enumerate(weeks):
        with week_cols[idx]:
            is_selected = (st.session_state.selected_week == idx + 1)

            if st.button(
                f"{week_info['emoji']} {week_info['label']}\n{week_info['desc']}",
                key=f"week_{idx+1}",
                use_container_width=True
            ):
                st.session_state.selected_week = idx + 1
                st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)

    # 获取当前周的任务
    week_key = f"week_{st.session_state.selected_week}"
    current_week_data = practice_data.get(week_key, {})

    if not current_week_data:
        st.warning("本周任务尚未开放")
        return

    # 本周简介 - 轻松风格
    st.markdown(f"""
<div style="background: #F8F9FA; padding: 1.5rem; border-radius: 12px; margin-bottom: 2rem; border-left: 4px solid #667eea;">
    <div style="font-size: 1.1rem; font-weight: 600; color: #2D3436; margin-bottom: 0.5rem;">
        {current_week_data["title"]}
    </div>
    <div style="font-size: 0.95rem; color: #636E72; line-height: 1.6; margin-bottom: 1rem;">
        {current_week_data["objective"]}
    </div>
    <div style="font-size: 0.85rem; color: #667eea; font-style: italic;">
        💡 {current_week_data["focus"]}
    </div>
</div>
""", unsafe_allow_html=True)

    # 显示每日任务 - 卡片式，自由浏览
    # 第三周有path_a和path_b两种路径，默认选择path_a
    daily_tasks = current_week_data.get("daily_tasks", [])
    if not daily_tasks and "path_a" in current_week_data:
        # 如果有路径选择，默认使用path_a
        path_data = current_week_data.get("path_a", {})
        daily_tasks = path_data.get("daily_tasks", [])
        if daily_tasks:
            st.info(f"📍 当前路径：{path_data.get('title', '路径A')} - {path_data.get('condition', '')}")

    # 计算本周的起始和结束天数
    # 第4周特殊处理：显示第22-30天（9天），其他周显示7天
    if st.session_state.selected_week == 4:
        week_start = 22
        week_end = 30
    else:
        week_start = (st.session_state.selected_week - 1) * 7 + 1
        week_end = st.session_state.selected_week * 7

    week_tasks = [task for task in daily_tasks if week_start <= task["day"] <= week_end]

    for task in week_tasks:
        task_day = task["day"]
        day_str = str(task_day)

        # 获取当天3个习惯的完成状态
        day_completion = habits_completion.get(day_str, {})
        completed_count_today = sum(day_completion.values())
        all_completed = (completed_count_today == 3)

        if all_completed:
            card_style = """
                style="
                    background: linear-gradient(135deg, #E8F5E9 0%, #C8E6C9 100%);
                    border: none;
                    border-radius: 16px;
                    padding: 1.8rem;
                    margin-bottom: 1.5rem;
                    box-shadow: 0 8px 24px rgba(76, 175, 80, 0.15), 0 2px 8px rgba(76, 175, 80, 0.1);
                    position: relative;
                    overflow: hidden;
                "
            """
        else:
            card_style = """
                style="
                    background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%);
                    border: 1px solid rgba(102, 126, 234, 0.1);
                    border-radius: 16px;
                    padding: 1.8rem;
                    margin-bottom: 1.5rem;
                    box-shadow: 0 10px 30px rgba(0, 0, 0, 0.08), 0 4px 12px rgba(0, 0, 0, 0.05);
                    position: relative;
                    overflow: hidden;
                "
            """

        st.markdown(f"<div {card_style}>", unsafe_allow_html=True)

        # 任务头部
        col_left, col_right = st.columns([3, 2])

        with col_left:
            # 标题
            status_emoji = "✅" if all_completed else "📌"
            title_style = "color: #2D3436;" if not all_completed else "color: #4CAF50;"
            st.markdown(f"<div style='font-size: 1.1rem; font-weight: 600; {title_style} margin-bottom: 0.5rem;'>{status_emoji} 第{task_day}天：{task['task']}</div>", unsafe_allow_html=True)

            # 说明
            st.markdown(f"<div style='font-size: 0.9rem; color: #636E72; line-height: 1.6; margin-bottom: 1rem;'>{task['instruction']}</div>", unsafe_allow_html=True)

            # 如果有行动项，用更自然的方式展示
            if task.get("action_items"):
                with st.expander("📝 具体怎么做", expanded=False):
                    for item in task["action_items"]:
                        st.markdown(f"<div style='color: #636E72; margin: 0.5rem 0;'>• {item}</div>", unsafe_allow_html=True)

            # 如果有示例
            if task.get("examples"):
                with st.expander("💡 参考一下", expanded=False):
                    for key, value in task["examples"].items():
                        st.markdown(f"**{key}**: {value}")

        with col_right:
            # 三个习惯的打卡checkbox
            st.markdown("**今日打卡**", unsafe_allow_html=True)
            st.markdown(f"<small>已完成 {completed_count_today}/3 个习惯</small>", unsafe_allow_html=True)

            for habit_idx, habit in enumerate(user_habits[:3]):
                habit_key = f"day_{task_day}_habit_{habit_idx}"
                is_habit_completed = day_completion.get(habit, False)

                new_status = st.checkbox(
                    habit,
                    key=habit_key,
                    value=is_habit_completed,
                    label_visibility="visible"
                )

                # 只有状态变化时才更新
                if new_status != is_habit_completed:
                    # 确保数据结构存在
                    if "habits_completion" not in st.session_state.practice_tracker[book_title]:
                        st.session_state.practice_tracker[book_title]["habits_completion"] = {}
                    if day_str not in st.session_state.practice_tracker[book_title]["habits_completion"]:
                        st.session_state.practice_tracker[book_title]["habits_completion"][day_str] = {}

                    st.session_state.practice_tracker[book_title]["habits_completion"][day_str][habit] = new_status

                    # 检查成就（打卡时）
                    check_and_unlock_achievements()

                    st.rerun()

        st.markdown("</div>", unsafe_allow_html=True)

    # 成就展示 - 只在有成就时显示
    if tracker.get("badges"):
        st.markdown('<div style="margin: 2rem 0; padding: 1.5rem; background: linear-gradient(135deg, #FFF9E6 0%, #FFF3CD 100%); border-radius: 16px; text-align: center;">', unsafe_allow_html=True)
        st.markdown('<div style="font-size: 1rem; color: #856404; margin-bottom: 1rem;">🏆 解锁的成就</div>', unsafe_allow_html=True)

        badges_html = '<div style="display: flex; gap: 0.75rem; justify-content: center; flex-wrap: wrap;">'
        for badge in tracker["badges"]:
            badges_html += f'<div style="background: linear-gradient(135deg, #FFD700 0%, #FFA500 100%); color: #ffffff; padding: 0.5rem 1.25rem; border-radius: 20px; font-weight: 600; font-size: 0.9rem; box-shadow: 0 4px 12px rgba(255, 215, 0, 0.3);">{badge}</div>'
        badges_html += '</div>'
        st.markdown(badges_html, unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)

    # 实用小贴士 - 使用 Streamlit 原生组件确保对齐
    st.markdown('<div style="margin: 2rem 0;">', unsafe_allow_html=True)
    st.markdown('<div style="font-size: 1.1rem; font-weight: 600; color: #2D3436; margin-bottom: 1.5rem; text-align: center;">💡 实用小贴士</div>', unsafe_allow_html=True)

    # 使用 expander 展示提示
    tips = [
        {
            "emoji": "📝",
            "title": "每天记录",
            "content": "花1分钟记录今天的实践感受，比如：「今天我按时完成了2分钟阅读，感觉很不错！」"
        },
        {
            "emoji": "🔄",
            "title": "定期回顾",
            "content": "每周日花10分钟回顾这周的完成情况，调整下周计划"
        },
        {
            "emoji": "🎯",
            "title": "关注连续",
            "content": "连续打卡会解锁成就徽章，让习惯养成更有趣"
        }
    ]

    for tip in tips:
        with st.expander(f"{tip['emoji']} {tip['title']}", expanded=False):
            st.markdown(f"<div style='color: #636E72; line-height: 1.6; padding: 0.5rem 0;'>{tip['content']}</div>", unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

    # 底部快速导航
    st.markdown('<div style="margin: 3rem 0 2rem 0;">', unsafe_allow_html=True)
    st.markdown("---")
    st.markdown('<div style="font-size: 1.1rem; font-weight: 600; color: #2D3436; margin: 2rem 0 1rem 0; text-align: center;">📅 快速切换周</div>', unsafe_allow_html=True)

    nav_cols = st.columns(4)
    for idx in range(4):
        with nav_cols[idx]:
            week_num = idx + 1
            is_current = (st.session_state.selected_week == week_num)
            emoji = ["🌱", "🌿", "🌳", "🏆"][idx]

            if st.button(
                f"{emoji} 第{week_num}周",
                key=f"nav_week_{week_num}",
                use_container_width=True
            ):
                st.session_state.selected_week = week_num
                st.rerun()

    st.markdown('</div>', unsafe_allow_html=True)


def calculate_consecutive_days(completed_days):
    """计算连续完成天数"""
    if not completed_days:
        return 0

    consecutive = 0
    for day in range(1, 32):  # 最多31天
        if str(day) in completed_days and completed_days[str(day)]:
            consecutive += 1
        else:
            break

    return consecutive


def render_reflection(content):
    """反思页 - 优化版"""
    reflection = content["reflection"]

    # 更新阅读统计
    current_book = st.session_state.current_book
    if current_book and current_book not in st.session_state.reading_stats["total_books_read"]:
        st.session_state.reading_stats["total_books_read"].add(current_book)

        # 安排复习计划（仅在新完成阅读时）
        book_content = st.session_state.current_content
        if book_content:
            schedule_review(current_book, book_content["title"])

    # 计算本次阅读时长并累加
    if 'reading_start_time' in st.session_state:
        import time
        elapsed = time.time() - st.session_state.reading_start_time
        st.session_state.reading_stats["total_reading_time"] += int(elapsed)
        st.session_state.reading_stats["last_read_time"] = time.time()

    # 触发成就检查（完成书籍时）
    check_and_unlock_achievements()

    # 滚动到顶部
    scroll_to_top()

    # 顶部导航
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("← 实践", key="reflection_back_practice"):
            st.session_state.page_rerun += 1
            st.session_state.current_section = "practice"
            st.rerun()

    st.markdown(f'<div class="section-title">{reflection["title"]}</div>', unsafe_allow_html=True)
    st.markdown(f'<div class="section-subtitle">{reflection["subtitle"]}</div>', unsafe_allow_html=True)

    # 阅读进度条
    render_progress_bar("reflection")

    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    # 思考题 - 优化布局
    for idx, question in enumerate(reflection["questions"], 1):
        st.markdown(f'<div class="section-block">', unsafe_allow_html=True)

        # 优雅的问题卡片设计 - 使用单行HTML格式
        question_html = f"""
<div style="background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%); border-radius: 16px; padding: 1.5rem; border: 1px solid rgba(102, 126, 234, 0.1); box-shadow: 0 2px 8px rgba(0, 0, 0, 0.04); margin-bottom: 1rem;">
    <div style="display: flex; gap: 1rem; align-items: flex-start;">
        <div style="background: linear-gradient(145deg, #667eea 0%, #764ba2 100%); color: #ffffff; width: 36px; height: 36px; border-radius: 50%; display: flex; align-items: center; justify-content: center; font-weight: 700; font-size: 1rem; flex-shrink: 0;">
            {idx}
        </div>
        <div style="flex: 1;">
            <div style="font-size: 1.05rem; color: #2D3436; line-height: 1.7; font-weight: 500; margin-bottom: 0.75rem;">
                {question["text"]}
            </div>
            <div style="font-size: 0.85rem; color: #636E72; line-height: 1.6; font-style: italic;">
                💡 {question["hint"]}
            </div>
        </div>
    </div>
</div>
"""
        st.markdown(question_html, unsafe_allow_html=True)

        # 笔记输入框
        st.markdown('<div style="margin-top: 1rem;">', unsafe_allow_html=True)
        user_note = st.text_area(
            "笔记",
            key=f"note_{idx}",
            placeholder="在这里记录你的思考，让想法更深刻...",
            height=100,
            label_visibility="collapsed"
        )
        st.markdown('</div>', unsafe_allow_html=True)

        if user_note:
            st.success("✓ 已记录")
            st.session_state.notes[f"q{idx}"] = user_note

            # 触发成就检查（记录反思时）
            check_and_unlock_achievements()

        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)

    # 金句回顾 - 优化设计
    st.markdown('<div class="section-block">', unsafe_allow_html=True)
    st.markdown('<div class="subsection-header">值得记住的话</div>', unsafe_allow_html=True)

    for quote in content["quotes"]:
        # 使用单行HTML格式
        quote_html = f"""
<div style="background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%); border-left: 4px solid #667eea; padding: 1.5rem; margin-bottom: 1rem; border-radius: 0 12px 12px 0; box-shadow: 0 2px 8px rgba(0, 0, 0, 0.04);">
    <div style="font-size: 1.1rem; line-height: 1.8; color: #2D3436; font-style: italic; position: relative;">
        {quote}
    </div>
</div>
"""
        st.markdown(quote_html, unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

    # ========== 分享功能区域 ==========
    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">✨ 分享你的阅读</div>', unsafe_allow_html=True)
    st.markdown('<div style="text-align: center; color: #636E72; font-size: 0.85rem; margin-bottom: 2rem;">生成精美卡片，分享到朋友圈或社交媒体</div>', unsafe_allow_html=True)

    # 金句选择
    selected_quote = st.selectbox(
        "选择要分享的金句",
        content["quotes"],
        key="share_quote_select"
    )

    # 按钮组 - 垂直布局
    if st.button("🎨 金句卡片", key=f"quote_card_{content['title']}", use_container_width=True):
        # 只处理换行
        quote_display = selected_quote.replace('\n', '<br/>')
        title_display = content['title'].replace('\n', '<br/>')
        author_display = content['author'].replace('\n', '<br/>')

        # HTML 卡片（用于截图）
        card_html = f'<div style="width: 100%; max-width: 500px; margin: 2rem auto; padding: 3rem 2rem; background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%); border-radius: 20px; box-shadow: 0 10px 40px rgba(0, 0, 0, 0.08); text-align: center; position: relative; overflow: hidden; border: 1px solid rgba(102, 126, 234, 0.1);"><div style="position: absolute; top: 0; left: 0; right: 0; height: 4px; background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);"></div><div style="margin-bottom: 2.5rem;"><div style="font-size: 1.1rem; color: #667eea; font-weight: 600; margin-bottom: 0.5rem;">{title_display}</div><div style="font-size: 0.9rem; color: #636E72; font-style: italic;">{author_display}</div></div><div style="background: linear-gradient(145deg, #f8f9fa 0%, #e8eef2 100%); border-radius: 16px; padding: 2rem; margin-bottom: 2.5rem; border: 1px solid rgba(102, 126, 234, 0.1);"><div style="font-size: 1.4rem; line-height: 1.9; color: #2D3436; font-weight: 600; position: relative; display: inline-block;">{quote_display}</div></div><div style="display: flex; flex-direction: column; gap: 0.75rem; align-items: center;"><div style="display: flex; align-items: center; gap: 0.75rem; padding: 0.75rem 1.5rem; background: rgba(102, 126, 234, 0.1); border-radius: 25px;"><span style="font-size: 1.2rem;">🧠</span><span style="color: #667eea; font-weight: 600; font-size: 0.95rem;">DeepRead 深读</span></div><div style="font-size: 0.75rem; color: #636E72; font-style: italic;">深度阅读 · 沉浸思考</div></div></div>'

        st.markdown(card_html, unsafe_allow_html=True)

        # 提示信息
        st.markdown('<div style="text-align: center; color: #636E72; font-size: 0.85rem; margin: 1rem 0;">💡 提示：如果下载的图片中文显示不正确，请直接截图上方卡片</div>', unsafe_allow_html=True)

        # 下载图片按钮
        img_data = create_quote_card_image(content['title'], content['author'], selected_quote)
        st.download_button(
            label="⬇️ 下载金句卡片图片",
            data=img_data,
            file_name=f"金句卡片_{content['title']}.png",
            mime="image/png",
            use_container_width=True,
            key=f"download_card_{content['title']}"
        )

    # 阅读海报生成
    st.markdown('<div style="margin-top: 2rem;"></div>', unsafe_allow_html=True)

    if st.button("📊 阅读海报", key="generate_poster", use_container_width=True):
        # 获取书籍信息
        book_info = next((b for b in BOOKS_DATA if b['title'] == content['title']), None)

        # 计算阅读统计
        stats = st.session_state.reading_stats
        books_read = len(stats["total_books_read"])
        total_hours = stats["total_reading_time"] // 3600
        total_minutes = (stats["total_reading_time"] % 3600) // 60

        if total_hours > 0:
            time_display = f"{total_hours}小时{total_minutes}分钟"
        elif total_minutes > 0:
            time_display = f"{total_minutes}分钟"
        else:
            time_display = "刚刚开始"

        # 生成阅读海报HTML
        quote_display = selected_quote.replace('\n', '<br/>')
        title_display = content['title'].replace('\n', '<br/>')
        author_display = content['author'].replace('\n', '<br/>')
        tags_html = "".join([f'<span style="background: rgba(102, 126, 234, 0.1); color: #667eea; padding: 0.3rem 0.8rem; border-radius: 15px; font-size: 0.75rem;">{tag}</span>' for tag in book_info['tags'] if book_info])
        poster_html = f'''<div style="width: 100%; max-width: 500px; margin: 2rem auto; padding: 0; background: linear-gradient(145deg, #ffffff 0%, #f8f9fa 100%); border-radius: 20px; box-shadow: 0 10px 40px rgba(0, 0, 0, 0.08); overflow: hidden; border: 1px solid rgba(102, 126, 234, 0.1);">
<div style="padding: 3rem 2rem 2rem 2rem; text-align: center; position: relative;">
    <div style="font-size: 5rem; margin-bottom: 1rem;">{book_info['emoji'] if book_info else '📖'}</div>
    <div style="font-size: 1.6rem; color: #2D3436; font-weight: 700; margin-bottom: 0.5rem;">{title_display}</div>
    <div style="font-size: 1rem; color: #636E72; font-weight: 400; margin-bottom: 1rem;">{author_display}</div>
    <div style="display: flex; gap: 0.5rem; justify-content: center; flex-wrap: wrap;">{tags_html}</div>
</div>
<div style="padding: 2rem;">
    <div style="background: #F8F9FA; border-left: 4px solid #667eea; padding: 1.5rem; border-radius: 8px; margin-bottom: 2rem;">
        <div style="font-size: 1.1rem; line-height: 1.8; color: #2D3436; font-style: italic;">{quote_display}</div>
    </div>
    <div style="display: flex; flex-direction: column; gap: 1.5rem; padding: 1.5rem 0;">
        <div style="display: flex; align-items: center; gap: 1rem; padding: 1rem; background: rgba(102, 126, 234, 0.05); border-radius: 12px;">
            <div style="background: linear-gradient(145deg, #667eea 0%, #764ba2 100%); color: #ffffff; width: 50px; height: 50px; border-radius: 12px; display: flex; align-items: center; justify-content: center; font-size: 1.5rem;">📚</div>
            <div style="flex: 1;">
                <div style="font-size: 2rem; color: #667eea; font-weight: 700;">{books_read}</div>
                <div style="font-size: 0.85rem; color: #636E72; margin-top: 0.25rem;">已读书籍</div>
            </div>
        </div>
        <div style="display: flex; align-items: center; gap: 1rem; padding: 1rem; background: rgba(118, 75, 162, 0.05); border-radius: 12px;">
            <div style="background: linear-gradient(145deg, #764ba2 0%, #667eea 100%); color: #ffffff; width: 50px; height: 50px; border-radius: 12px; display: flex; align-items: center; justify-content: center; font-size: 1.5rem;">⏱️</div>
            <div style="flex: 1;">
                <div style="font-size: 2rem; color: #764ba2; font-weight: 700;">{time_display}</div>
                <div style="font-size: 0.85rem; color: #636E72; margin-top: 0.25rem;">阅读时长</div>
            </div>
        </div>
    </div>
</div>
<div style="text-align: center; margin-top: 1.5rem;">
    <div style="display: inline-flex; align-items: center; gap: 0.5rem; padding: 0.75rem 1.5rem; background: rgba(102, 126, 234, 0.1); border-radius: 25px; color: #667eea; font-weight: 600; font-size: 0.9rem;">
        <span>🧠</span><span>DeepRead 深读</span>
    </div>
    <div style="font-size: 0.75rem; color: #636E72; margin-top: 0.75rem; font-style: italic;">深度阅读 · 沉浸思考</div>
</div>
</div>'''
        st.markdown(poster_html, unsafe_allow_html=True)

        # 提示信息
        st.markdown('<div style="text-align: center; color: #636E72; font-size: 0.85rem; margin: 1rem 0;">💡 提示：如果下载的图片中文显示不正确，请直接截图上方海报</div>', unsafe_allow_html=True)

        # 下载图片按钮
        poster_stats = {
            'books_read': books_read,
            'time_display': time_display
        }
        img_data = create_reading_poster_image(
            content['title'],
            content['author'],
            book_info['emoji'] if book_info else '📖',
            book_info['tags'] if book_info else [],
            selected_quote,
            poster_stats
        )
        st.download_button(
            label="⬇️ 下载阅读海报图片",
            data=img_data,
            file_name=f"阅读海报_{content['title']}.png",
            mime="image/png",
            use_container_width=True,
            key=f"download_poster_{content['title']}"
        )

    # 分享文案
    st.markdown('<div style="margin-top: 2rem;"></div>', unsafe_allow_html=True)

    if st.button("📋 分享文案", key="copy_share", use_container_width=True):
        # 生成分享文案
        share_text = f"""📚 {content['title']} - {content['author']}

💡 核心观点：
{selected_quote}

📖 我的思考：
"""
        # 添加用户笔记
        for key, value in st.session_state.notes.items():
            if value:
                share_text += f"\n{value}\n"

        share_text += f"""
🧠 来自 DeepRead 深读
深度阅读 · 沉浸思考

👉 一起读书成长吧！
"""

        st.markdown(f'<div style="background: #F8F9FA; border-left: 4px solid #667eea; padding: 1.5rem; border-radius: 8px; margin: 1rem 0;"><div style="font-size: 0.9rem; color: #636E72; margin-bottom: 0.75rem; font-weight: 600;">📋 分享文案（可复制）</div><div style="font-size: 0.85rem; line-height: 1.8; color: #2D3436; white-space: pre-wrap; font-family: \'Noto Serif SC\', serif; background: #ffffff; padding: 1rem; border-radius: 6px; border: 1px solid #E8EEF2;">{share_text}</div><div style="font-size: 0.8rem; color: #636E72; margin-top: 0.75rem; font-style: italic;">💡 复制上方文字，分享到朋友圈、微博、小红书等平台</div></div>', unsafe_allow_html=True)
    # ============================================

    # 导出功能区
    st.markdown('<div class="section-divider"></div>', unsafe_allow_html=True)
    st.markdown('<div class="section-title">📤 导出学习笔记</div>', unsafe_allow_html=True)

    st.markdown('<div style="text-align: center; color: #636E72; font-size: 0.85rem; margin-bottom: 2rem;">选择导出格式，保存你的阅读成果</div>', unsafe_allow_html=True)

    # 导出选项卡
    export_tab1, export_tab2, export_tab3 = st.columns(3)

    with export_tab1:
        st.markdown('<div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); padding: 1.5rem; border-radius: 12px; text-align: center; color: white; margin-bottom: 1rem;"><div style="font-size: 2rem; margin-bottom: 0.5rem;">📝</div><div style="font-size: 1rem; font-weight: 600; margin-bottom: 0.25rem;">我的笔记</div><div style="font-size: 0.75rem; opacity: 0.9;">仅导出个人思考</div></div>', unsafe_allow_html=True)

        if st.button("Markdown", key="export_notes_md", use_container_width=True):
            md_content = generate_notes_only(content, st.session_state.notes)
            filename = f"{content['title']}_我的笔记_{datetime.now().strftime('%Y%m%d')}.md"
            st.download_button(
                label="⬇️ 下载MD文件",
                data=md_content,
                file_name=filename,
                mime="text/markdown",
                key="download_notes_md"
            )

        # Word导出
        if WORD_SUPPORT:
            if st.button("Word文档", key="export_notes_word", use_container_width=True):
                word_bytes = generate_word_bytes(content, st.session_state.notes, include_full_content=False)
                if word_bytes:
                    filename = f"{content['title']}_我的笔记_{datetime.now().strftime('%Y%m%d')}.docx"
                    st.download_button(
                        label="⬇️ 下载Word文件",
                        data=word_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_notes_word"
                    )
        else:
            st.info("💡 安装python-docx库以支持Word导出")

        # PDF导出
        if PDF_SUPPORT:
            if st.button("PDF文档", key="export_notes_pdf", use_container_width=True):
                pdf_bytes = generate_pdf_bytes(content, st.session_state.notes, include_full_content=False)
                if pdf_bytes:
                    filename = f"{content['title']}_我的笔记_{datetime.now().strftime('%Y%m%d')}.pdf"
                    st.download_button(
                        label="⬇️ 下载PDF文件",
                        data=pdf_bytes,
                        file_name=filename,
                        mime="application/pdf",
                        key="download_notes_pdf"
                    )
        else:
            st.info("💡 安装reportlab库以支持PDF导出")

    with export_tab2:
        st.markdown('<div style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%); padding: 1.5rem; border-radius: 12px; text-align: center; color: white; margin-bottom: 1rem;"><div style="font-size: 2rem; margin-bottom: 0.5rem;">📚</div><div style="font-size: 1rem; font-weight: 600; margin-bottom: 0.25rem;">完整笔记</div><div style="font-size: 0.75rem; opacity: 0.9;">包含所有内容</div></div>', unsafe_allow_html=True)

        if st.button("Markdown", key="export_full_md", use_container_width=True):
            md_content = generate_markdown(content, st.session_state.notes)
            filename = f"{content['title']}_完整学习笔记_{datetime.now().strftime('%Y%m%d')}.md"
            st.download_button(
                label="⬇️ 下载MD文件",
                data=md_content,
                file_name=filename,
                mime="text/markdown",
                key="download_full_md"
            )

        # Word导出
        if WORD_SUPPORT:
            if st.button("Word文档", key="export_full_word", use_container_width=True):
                word_bytes = generate_word_bytes(content, st.session_state.notes, include_full_content=True)
                if word_bytes:
                    filename = f"{content['title']}_完整学习笔记_{datetime.now().strftime('%Y%m%d')}.docx"
                    st.download_button(
                        label="⬇️ 下载Word文件",
                        data=word_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key="download_full_word"
                    )
        else:
            st.info("💡 安装python-docx库")

        # PDF导出
        if PDF_SUPPORT:
            if st.button("PDF文档", key="export_full_pdf", use_container_width=True):
                pdf_bytes = generate_pdf_bytes(content, st.session_state.notes, include_full_content=True)
                if pdf_bytes:
                    filename = f"{content['title']}_完整学习笔记_{datetime.now().strftime('%Y%m%d')}.pdf"
                    st.download_button(
                        label="⬇️ 下载PDF文件",
                        data=pdf_bytes,
                        file_name=filename,
                        mime="application/pdf",
                        key="download_full_pdf"
                    )
        else:
            st.info("💡 安装reportlab库")

    with export_tab3:
        st.markdown('<div style="background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%); padding: 1.5rem; border-radius: 12px; text-align: center; color: white; margin-bottom: 1rem;"><div style="font-size: 2rem; margin-bottom: 0.5rem;">📊</div><div style="font-size: 1rem; font-weight: 600; margin-bottom: 0.25rem;">使用指南</div><div style="font-size: 0.75rem; opacity: 0.9;">导出说明</div></div>', unsafe_allow_html=True)

        st.markdown('<div style="background: #F8F9FA; padding: 1.5rem; border-radius: 12px; font-size: 0.85rem; line-height: 1.8;"><div style="margin-bottom: 1rem;"><strong>📝 Markdown (.md)</strong><br/>适合导入飞书、Notion等笔记软件</div><div style="margin-bottom: 1rem;"><strong>📄 Word (.docx)</strong><br/>适合编辑和分享，格式完整</div><div><strong>📕 PDF (.pdf)</strong><br/>适合打印和归档，格式固定</div></div>', unsafe_allow_html=True)

    # 完成阅读
    st.markdown('<div class="nav-container">', unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 2, 1])
    with col1:
        if st.button("← 返回", key="reflection_back"):
            st.session_state.page_rerun += 1
            st.session_state.current_section = "practice"
            st.rerun()

    with col2:
        if st.button("📚 返回书库", key="reflection_to_library", use_container_width=True):
            st.session_state.page_rerun += 1
            st.session_state.current_book = None
            st.session_state.current_content = None
            st.session_state.current_section = "library"
            st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)


def render_statistics():
    """显示详细统计页面"""
    st.markdown("---")
    st.markdown('<div class="section-title">📊 阅读数据统计</div>', unsafe_allow_html=True)

    # 试用状态卡片
    if st.session_state.user_tier == "trial":
        days_remaining = get_trial_days_remaining()
        if days_remaining > 0:
            st.info(f"🎁 **7天深度版免费试用中** - 还剩 {days_remaining} 天，升级后解锁更多功能")
        else:
            st.warning("⏰ **试用已到期** - 升级深度版继续使用数据统计功能")

    # 核心统计数据
    stats = st.session_state.reading_stats

    # 统计卡片（3列）
    col1, col2, col3 = st.columns(3)

    with col1:
        books_read_count = len(stats["total_books_read"])
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                    padding: 1.5rem; border-radius: 12px; text-align: center; color: white;">
            <div style="font-size: 0.85rem; opacity: 0.9; margin-bottom: 0.5rem;">已读书籍</div>
            <div style="font-size: 2.5rem; font-weight: 700;">{books_read_count}</div>
            <div style="font-size: 0.75rem; opacity: 0.8; margin-top: 0.5rem;">本</div>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        total_hours = stats["total_reading_time"] // 3600
        total_minutes = (stats["total_reading_time"] % 3600) // 60

        if total_hours > 0:
            time_display = f"{total_hours}h {total_minutes}m"
        elif total_minutes > 0:
            time_display = f"{total_minutes}m"
        else:
            time_display = "0m"

        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #f093fb 0%, #f5576c 100%);
                    padding: 1.5rem; border-radius: 12px; text-align: center; color: white;">
            <div style="font-size: 0.85rem; opacity: 0.9; margin-bottom: 0.5rem;">阅读时长</div>
            <div style="font-size: 2.5rem; font-weight: 700;">{time_display}</div>
            <div style="font-size: 0.75rem; opacity: 0.8; margin-top: 0.5rem;">总计</div>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        # 连续签到天数（这里用读书天数代替）
        reading_days = len(stats["daily_progress"])
        st.markdown(f"""
        <div style="background: linear-gradient(135deg, #4facfe 0%, #00f2fe 100%);
                    padding: 1.5rem; border-radius: 12px; text-align: center; color: white;">
            <div style="font-size: 0.85rem; opacity: 0.9; margin-bottom: 0.5rem;">活跃天数</div>
            <div style="font-size: 2.5rem; font-weight: 700;">{reading_days}</div>
            <div style="font-size: 0.75rem; opacity: 0.8; margin-top: 0.5rem;">天</div>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # 已读书籍列表
    if stats["total_books_read"]:
        st.markdown("### 📚 已完成阅读")

        for book_id in stats["total_books_read"]:
            book = next((b for b in st.session_state.books if b["id"] == book_id), None)
            if book:
                with st.expander(f"📖 {book['title']}"):
                    # 获取该书的阅读时间
                    book_time = stats.get("book_reading_time", {}).get(book_id, 0)
                    book_hours = book_time // 3600
                    book_minutes = (book_time % 3600) // 60

                    if book_hours > 0:
                        time_str = f"{book_hours}小时{book_minutes}分钟"
                    else:
                        time_str = f"{book_minutes}分钟"

                    st.markdown(f"**阅读时长：** {time_str}")
                    st.markdown(f"**完成时间：** {stats.get('book_completion_date', {}).get(book_id, '未知')}")

                    # 显示该书的所有实践记录
                    if st.session_state.practices.get(book_id):
                        st.markdown("**实践记录：**")
                        for idx, practice in enumerate(st.session_state.practices[book_id], 1):
                            st.markdown(f"{idx}. {practice.get('action', '')[:50]}...")
    else:
        st.info("📚 还没有完成阅读的书籍，继续加油！")

    st.markdown("<br>", unsafe_allow_html=True)

    # 成就系统（完整版）
    st.markdown("### 🏆 阅读成就")

    # 获取成就进度
    achievement_progress = get_achievement_progress()

    # 按等级分组展示成就
    tier1_achievements = [a for a in ACHIEVEMENTS_DEFINITIONS.values() if a["tier"] == 1]
    tier2_achievements = [a for a in ACHIEVEMENTS_DEFINITIONS.values() if a["tier"] == 2]
    tier3_achievements = [a for a in ACHIEVEMENTS_DEFINITIONS.values() if a["tier"] == 3]

    # Tier 1 成就（初级）
    if tier1_achievements:
        st.markdown("#### ⭐ 初级成就")
        tier1_cols = st.columns(min(len(tier1_achievements), 3))

        for i, achievement in enumerate(tier1_achievements):
            with tier1_cols[i % 3]:
                progress = achievement_progress[achievement["id"]]
                is_unlocked = progress["unlocked"]

                # 进度条颜色
                if is_unlocked:
                    progress_color = "#27ae60"
                    bg_color = "#d4edda"
                    border_color = "#28a745"
                else:
                    progress_color = "#667eea"
                    bg_color = "#f0f0f0"
                    border_color = "#ddd"

                st.markdown(f"""
                <div style="background: {bg_color};
                           padding: 1rem; border-radius: 10px;
                           border: 2px solid {border_color}; margin-bottom: 0.5rem;">
                    <div style="display: flex; align-items: center; gap: 0.5rem; margin-bottom: 0.5rem;">
                        <div style="font-size: 2rem;">{achievement['icon'] if is_unlocked else '🔒'}</div>
                        <div style="flex: 1;">
                            <div style="font-size: 0.9rem; font-weight: 600; margin-bottom: 0.25rem;">
                                {achievement['name']}
                            </div>
                            <div style="font-size: 0.7rem; color: #636E72;">
                                {achievement['description']}
                            </div>
                        </div>
                    </div>
                    <div style="margin-top: 0.5rem;">
                        <div style="display: flex; justify-content: space-between; font-size: 0.7rem; margin-bottom: 0.25rem;">
                            <span>进度</span>
                            <span style="color: {progress_color}; font-weight: 600;">
                                {progress['current']}/{progress['target']}
                                {' ✓ 已解锁' if is_unlocked else f' ({progress["percent"]}%)'}
                            </span>
                        </div>
                        <div style="background: #e0e0e0; height: 8px; border-radius: 4px; overflow: hidden;">
                            <div style="background: {progress_color}; height: 100%;
                                       width: {progress['percent']}%; transition: width 0.3s ease;"></div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # Tier 2 成就（中级）
    if tier2_achievements:
        st.markdown("#### 🌟 中级成就")
        tier2_cols = st.columns(min(len(tier2_achievements), 3))

        for i, achievement in enumerate(tier2_achievements):
            with tier2_cols[i % 3]:
                progress = achievement_progress[achievement["id"]]
                is_unlocked = progress["unlocked"]

                # 进度条颜色
                if is_unlocked:
                    progress_color = "#27ae60"
                    bg_color = "#d4edda"
                    border_color = "#28a745"
                else:
                    progress_color = "#f39c12"
                    bg_color = "#fef5e7"
                    border_color = "#f39c12"

                st.markdown(f"""
                <div style="background: {bg_color};
                           padding: 1rem; border-radius: 10px;
                           border: 2px solid {border_color}; margin-bottom: 0.5rem;">
                    <div style="display: flex; align-items: center; gap: 0.5rem; margin-bottom: 0.5rem;">
                        <div style="font-size: 2rem;">{achievement['icon'] if is_unlocked else '🔒'}</div>
                        <div style="flex: 1;">
                            <div style="font-size: 0.9rem; font-weight: 600; margin-bottom: 0.25rem;">
                                {achievement['name']}
                            </div>
                            <div style="font-size: 0.7rem; color: #636E72;">
                                {achievement['description']}
                            </div>
                        </div>
                    </div>
                    <div style="margin-top: 0.5rem;">
                        <div style="display: flex; justify-content: space-between; font-size: 0.7rem; margin-bottom: 0.25rem;">
                            <span>进度</span>
                            <span style="color: {progress_color}; font-weight: 600;">
                                {progress['current']}/{progress['target']}
                                {' ✓ 已解锁' if is_unlocked else f' ({progress["percent"]}%)'}
                            </span>
                        </div>
                        <div style="background: #e0e0e0; height: 8px; border-radius: 4px; overflow: hidden;">
                            <div style="background: {progress_color}; height: 100%;
                                       width: {progress['percent']}%; transition: width 0.3s ease;"></div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # Tier 3 成就（高级）
    if tier3_achievements:
        st.markdown("#### 👑 高级成就")
        tier3_cols = st.columns(min(len(tier3_achievements), 3))

        for i, achievement in enumerate(tier3_achievements):
            with tier3_cols[i % 3]:
                progress = achievement_progress[achievement["id"]]
                is_unlocked = progress["unlocked"]

                # 进度条颜色
                if is_unlocked:
                    progress_color = "#27ae60"
                    bg_color = "#d4edda"
                    border_color = "#28a745"
                else:
                    progress_color = "#9b59b6"
                    bg_color = "#f4ecf7"
                    border_color = "#9b59b6"

                st.markdown(f"""
                <div style="background: {bg_color};
                           padding: 1rem; border-radius: 10px;
                           border: 2px solid {border_color}; margin-bottom: 0.5rem;">
                    <div style="display: flex; align-items: center; gap: 0.5rem; margin-bottom: 0.5rem;">
                        <div style="font-size: 2rem;">{achievement['icon'] if is_unlocked else '🔒'}</div>
                        <div style="flex: 1;">
                            <div style="font-size: 0.9rem; font-weight: 600; margin-bottom: 0.25rem;">
                                {achievement['name']}
                            </div>
                            <div style="font-size: 0.7rem; color: #636E72;">
                                {achievement['description']}
                            </div>
                        </div>
                    </div>
                    <div style="margin-top: 0.5rem;">
                        <div style="display: flex; justify-content: space-between; font-size: 0.7rem; margin-bottom: 0.25rem;">
                            <span>进度</span>
                            <span style="color: {progress_color}; font-weight: 600;">
                                {progress['current']}/{progress['target']}
                                {' ✓ 已解锁' if is_unlocked else f' ({progress["percent"]}%)'}
                            </span>
                        </div>
                        <div style="background: #e0e0e0; height: 8px; border-radius: 4px; overflow: hidden;">
                            <div style="background: {progress_color}; height: 100%;
                                       width: {progress['percent']}%; transition: width 0.3s ease;"></div>
                        </div>
                    </div>
                </div>
                """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)

    # 功能对比表（仅在免费版显示）
    if st.session_state.user_tier != "premium":
        st.markdown("### ✨ 功能对比")

        st.markdown("""
        <div style="overflow-x: auto;">
            <table style="width: 100%; border-collapse: collapse; margin: 1rem 0;">
                <thead>
                    <tr style="background: #f0f0f0;">
                        <th style="padding: 0.75rem; text-align: left; border: 1px solid #ddd;">功能</th>
                        <th style="padding: 0.75rem; text-align: center; border: 1px solid #ddd;">免费版</th>
                        <th style="padding: 0.75rem; text-align: center; border: 1px solid #ddd; background: #e3f2fd;">深度版</th>
                    </tr>
                </thead>
                <tbody>
                    <tr>
                        <td style="padding: 0.75rem; border: 1px solid #ddd;">📖 完整阅读内容</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd;">✅</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd; background: #e3f2fd;">✅</td>
                    </tr>
                    <tr>
                        <td style="padding: 0.75rem; border: 1px solid #ddd;">✍️ 基础实践笔记</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd;">✅</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd; background: #e3f2fd;">✅</td>
                    </tr>
                    <tr>
                        <td style="padding: 0.75rem; border: 1px solid #ddd;">💾 本地数据存储</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd;">✅</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd; background: #e3f2fd;">✅</td>
                    </tr>
                    <tr>
                        <td style="padding: 0.75rem; border: 1px solid #ddd;">📊 阅读数据统计</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd;">7天</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd; background: #e3f2fd;">✅ 永久</td>
                    </tr>
                    <tr>
                        <td style="padding: 0.75rem; border: 1px solid #ddd;">☁️ 云端数据同步</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd;">❌</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd; background: #e3f2fd;">✅</td>
                    </tr>
                    <tr>
                        <td style="padding: 0.75rem; border: 1px solid #ddd;">🧠 智能复习提醒</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd;">❌</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd; background: #e3f2fd;">✅</td>
                    </tr>
                    <tr>
                        <td style="padding: 0.75rem; border: 1px solid #ddd;">📤 高级导出功能</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd;">❌</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd; background: #e3f2fd;">✅</td>
                    </tr>
                    <tr>
                        <td style="padding: 0.75rem; border: 1px solid #ddd;">🤖 AI智能推荐</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd;">❌</td>
                        <td style="padding: 0.75rem; text-align: center; border: 1px solid #ddd; background: #e3f2fd;">✅</td>
                    </tr>
                </tbody>
            </table>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # 价格方案
        st.markdown("### 💎 选择适合你的方案")

        pricing_col1, pricing_col2, pricing_col3 = st.columns(3)

        with pricing_col1:
            st.markdown("""
            <div style="background: white; padding: 1.5rem; border-radius: 12px;
                        border: 2px solid #ddd; text-align: center; height: 100%;">
                <div style="font-size: 0.9rem; font-weight: 600; margin-bottom: 0.5rem; color: #636E72;">
                    月付
                </div>
                <div style="font-size: 2rem; font-weight: 700; color: #2D3436; margin-bottom: 0.5rem;">
                    ¥9.9
                </div>
                <div style="font-size: 0.75rem; color: #636E72; margin-bottom: 1rem;">
                    /月
                </div>
                <div style="font-size: 0.7rem; color: #999;">灵活订阅</div>
            </div>
            """, unsafe_allow_html=True)

        with pricing_col2:
            st.markdown("""
            <div style="background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
                        padding: 1.5rem; border-radius: 12px; text-align: center;
                        color: white; height: 100%; position: relative;">
                <div style="position: absolute; top: -10px; left: 50%; transform: translateX(-50%);
                            background: #fdcb6e; color: #2D3436; padding: 0.25rem 0.75rem;
                            border-radius: 10px; font-size: 0.7rem; font-weight: 600;">
                    推荐
                </div>
                <div style="font-size: 0.9rem; font-weight: 600; margin-bottom: 0.5rem; opacity: 0.9;">
                    季付
                </div>
                <div style="font-size: 2rem; font-weight: 700; margin-bottom: 0.5rem;">
                    ¥19.9
                </div>
                <div style="font-size: 0.75rem; opacity: 0.8; margin-bottom: 1rem;">
                    /季
                </div>
                <div style="font-size: 0.7rem; opacity: 0.9;">省 ¥9.8</div>
            </div>
            """, unsafe_allow_html=True)

        with pricing_col3:
            st.markdown("""
            <div style="background: white; padding: 1.5rem; border-radius: 12px;
                        border: 2px solid #ddd; text-align: center; height: 100%;">
                <div style="font-size: 0.9rem; font-weight: 600; margin-bottom: 0.5rem; color: #636E72;">
                    年付
                </div>
                <div style="font-size: 2rem; font-weight: 700; color: #2D3436; margin-bottom: 0.5rem;">
                    ¥59.9
                </div>
                <div style="font-size: 0.75rem; color: #636E72; margin-bottom: 1rem;">
                    /年
                </div>
                <div style="font-size: 0.7rem; color: #27ae60;">省 ¥58.9</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)

        # 升级按钮
        st.markdown("<div style='text-align: center; margin: 2rem 0;'>", unsafe_allow_html=True)
        if st.button("✨ 升级到深度版", use_container_width=True, key="upgrade_premium"):
            st.info("💡 升级功能即将开放，敬请期待！")
        st.markdown("</div>", unsafe_allow_html=True)

    # 返回按钮
    st.markdown("<br>", unsafe_allow_html=True)
    col1, col2, col3 = st.columns([1, 2, 1])

    with col2:
        if st.button("📚 返回书库", use_container_width=True, key="stats_back_to_library"):
            st.session_state.page_rerun += 1
            st.session_state.current_book = None
            st.session_state.current_section = "library"
            st.rerun()


def render_sidebar():
    """侧边栏 - 简化版（大号emoji）"""
    with st.sidebar:
        # Logo区域 - 大号emoji
        st.markdown("""
<div style="text-align: center; padding: 2.5rem 0 1.5rem 0; border-bottom: 2px solid #E8EEF2;">
    <div style="font-size: 4rem; margin-bottom: 0.75rem;">🧠</div>
    <div style="font-size: 1.5rem; font-weight: 600; color: #2D3436; margin-bottom: 0.5rem; letter-spacing: 0.05em; font-family: 'Noto Serif SC', serif;">DeepRead</div>
    <div style="font-size: 0.85rem; color: #636E72; margin-top: 0.75rem; font-style: italic; letter-spacing: 0.03em;">深度阅读 · 沉浸思考</div>
</div>
""", unsafe_allow_html=True)

        # ========== 阅读统计面板 ==========
        stats = st.session_state.reading_stats
        books_read_count = len(stats["total_books_read"])

        # 计算总阅读时长显示
        total_hours = stats["total_reading_time"] // 3600
        total_minutes = (stats["total_reading_time"] % 3600) // 60

        if total_hours > 0:
            time_display = f"{total_hours}小时{total_minutes}分钟"
        elif total_minutes > 0:
            time_display = f"{total_minutes}分钟"
        else:
            time_display = "0分钟"

        st.markdown('<div style="margin: 1.5rem 0 0.75rem 0;">', unsafe_allow_html=True)
        st.markdown('<div style="font-size: 0.75rem; font-weight: 600; color: #636E72; margin-bottom: 0.75rem;">📊 阅读统计</div>', unsafe_allow_html=True)

        st.markdown(f"""
<div style="background: #F0F3F5; padding: 0.875rem; border-radius: 8px; margin-bottom: 0.5rem;">
    <div style="font-size: 0.7rem; color: #636E72; margin-bottom: 0.25rem;">已读书籍</div>
    <div style="font-size: 0.9rem; font-weight: 600; color: #2D3436;">📚 {books_read_count} 本</div>
</div>

<div style="background: #F0F3F5; padding: 0.875rem; border-radius: 8px;">
    <div style="font-size: 0.7rem; color: #636E72; margin-bottom: 0.25rem;">累计阅读时长</div>
    <div style="font-size: 0.9rem; font-weight: 600; color: #2D3436;">⏱️ {time_display}</div>
</div>
""", unsafe_allow_html=True)

        # 详细统计按钮
        if st.button("📊 详细统计", key="nav_statistics", use_container_width=True):
            st.session_state.current_section = "statistics"
            st.session_state.page_rerun += 1
            st.rerun()
        # ==========================================

        # ========== 新功能：收藏书籍 ==========
        if 'favorite_books' not in st.session_state:
            st.session_state.favorite_books = []

        if st.session_state.favorite_books:
            st.markdown('<div style="margin: 1.5rem 0 0.75rem 0;">', unsafe_allow_html=True)
            st.markdown('<div style="font-size: 0.75rem; font-weight: 600; color: #636E72; margin-bottom: 0.75rem;">❤️ 我的收藏</div>', unsafe_allow_html=True)

            for fav_book in st.session_state.favorite_books:
                # 获取书籍信息
                book_info = next((b for b in BOOKS_DATA if b['title'] == fav_book), None)
                if book_info:
                    if st.button(f"📖 {fav_book}", key=f"sidebar_fav_{fav_book}", use_container_width=True):
                        st.session_state.page_rerun += 1
                        st.session_state.current_book = fav_book
                        st.session_state.current_content = get_book_content(fav_book)
                        st.session_state.current_section = "intro"
                        st.rerun()
        # ==========================================

        if st.session_state.current_book:
            # 当前阅读
            st.markdown(f"""
<div style="background: #F0F3F5; padding: 0.875rem; border-radius: 8px; margin: 1.25rem 0;">
    <div style="font-size: 0.7rem; color: #636E72; margin-bottom: 0.25rem;">正在阅读</div>
    <div style="font-size: 0.9rem; font-weight: 600; color: #2D3436;">{st.session_state.current_book}</div>
</div>
""", unsafe_allow_html=True)

            # ========== 新功能：阅读时长统计 ==========
            # 记录开始时间
            if 'reading_start_time' not in st.session_state:
                import time
                st.session_state.reading_start_time = time.time()

            # 计算阅读时长
            import time
            elapsed = time.time() - st.session_state.reading_start_time
            minutes = int(elapsed // 60)
            seconds = int(elapsed % 60)

            # 显示阅读时长
            time_text = f"{minutes}分{seconds}秒" if minutes > 0 else f"{seconds}秒"
            st.markdown(f"""
<div style="background: #F0F3F5; padding: 0.875rem; border-radius: 8px; margin: 1.25rem 0 0.75rem 0;">
    <div style="font-size: 0.7rem; color: #636E72; margin-bottom: 0.25rem;">本次阅读时长</div>
    <div style="font-size: 0.9rem; font-weight: 600; color: #2D3436;">⏱️ {time_text}</div>
</div>
""", unsafe_allow_html=True)
            # ==========================================

            # 阅读进度
            st.markdown('<div style="margin: 1.5rem 0 0.75rem 0;">', unsafe_allow_html=True)
            st.markdown('<div style="font-size: 0.75rem; font-weight: 600; color: #636E72; margin-bottom: 0.75rem;">阅读进度</div>', unsafe_allow_html=True)

            sections = [
                ("intro", "📖 引言"),
                ("insights", "💡 洞察"),
                ("practice", "✅ 实践"),
                ("reflection", "🤔 反思")
            ]

            for key, label in sections:
                is_current = st.session_state.current_section == key
                if is_current:
                    st.markdown(f"""
<div style="background: #636E72; color: #FFFFFF; padding: 0.65rem 0.875rem; border-radius: 6px; margin-bottom: 0.5rem; font-size: 0.85rem;">
    {label}
</div>
""", unsafe_allow_html=True)
                else:
                    if st.button(label, key=f"nav_{key}"):
                        st.session_state.page_rerun += 1
                        st.session_state.current_section = key
                        st.rerun()

            # 30天实践计划入口（如果书籍有）
            if st.session_state.current_book in PRACTICE_TASKS:
                st.markdown('<div style="margin: 1.5rem 0 0.75rem 0;">', unsafe_allow_html=True)

                tracker = st.session_state.practice_tracker.get(st.session_state.current_book, {})
                if tracker:
                    # 已开始，显示进度
                    current_day = tracker.get("current_day", 1)
                    completed_days = len([d for d, completed in tracker.get("completed_days", {}).items() if completed])

                    if st.button(f"🎯 实践计划 ({completed_days}/30)", key="nav_practice_tasks", use_container_width=True):
                        st.session_state.page_rerun += 1
                        st.session_state.current_section = "practice_tasks"
                        st.rerun()
                else:
                    # 未开始，显示开始按钮
                    if st.button("🎯 30天实践计划", key="nav_practice_tasks", use_container_width=True):
                        st.session_state.page_rerun += 1
                        st.session_state.current_section = "practice_tasks"
                        st.rerun()

                st.markdown('</div>', unsafe_allow_html=True)

            # 返回按钮
            st.markdown('<div style="margin-top: 1.5rem;">', unsafe_allow_html=True)
            if st.button("📚 返回书库", use_container_width=True):
                st.session_state.page_rerun += 1
                st.session_state.current_book = None
                st.session_state.current_content = None
                st.session_state.current_section = "library"
                st.rerun()

        # 底部信息 - 放大文字
        st.markdown("""
<div style="margin-top: auto; padding-top: 2rem; text-align: center; border-top: 2px solid #E8EEF2;">
    <div style="font-size: 0.85rem; color: #636E72; line-height: 1.8; font-weight: 500;">
        给自己时间<br/>慢慢来<br/><br/>🌱
    </div>
</div>
""", unsafe_allow_html=True)


def main():
    """主函数"""
    init_session_state()

    # 显示欢迎页（首次访问）
    if st.session_state.is_first_visit:
        show_welcome_page()
        return  # 欢迎页后直接返回，等待用户点击按钮

    # 显示试用提醒横幅
    if st.session_state.user_tier == "trial":
        show_trial_notice()

    # 显示成就解锁通知
    show_achievement_notifications()

    # 显示复习提醒面板
    show_review_reminder_panel()

    # 显示新手引导气泡
    show_guide_bubble()

    # 侧边栏
    render_sidebar()

    # 主内容区
    if not st.session_state.current_book:
        render_library()
    else:
        content = st.session_state.current_content
        section = st.session_state.current_section

        if section == "statistics":
            render_statistics()
        elif section == "intro":
            render_introduction(content)
        elif section == "insights":
            render_insights(content)
        elif section == "practice":
            render_practice(content)
        elif section == "practice_tasks":
            render_practice_tasks(content)
        elif section == "reflection":
            render_reflection(content)


if __name__ == "__main__":
    main()
